import dearpygui.dearpygui as dpg

import os 
import shutil

from shutil import copy

import string 
from string import ascii_letters

from docx import Document
from docx.shared import RGBColor

from openpyxl import Workbook, load_workbook
from openpyxl.styles import numbers

#Misc Variables
newline = "\n"
tab = "\t"

value_start_indicator = "{"
value_end_indicator = "}"
    
MAIN_FILE_PATH = os.path.abspath(__file__)
MAIN_FILE_NAME = str(os.path.basename(__file__)) 

backup_script = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "\\backup\\backup.py")
update_script = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "\\update\\update.py")
main_script = MAIN_FILE_PATH
#Misc Variables


#Classes
class DocType:
    def __init__(self):
        self.obj_name = []
        self.doc_type_name = []
        self.docs = []
class Doc:
    def __init__(self):
        self.obj_name = []
        self.doc_name = []
        self.doc_type = []
        self.doc_path = []    
        self.doc_objs = []
class Lawyer:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
        self.suffix = []
        self.bar_number = []
        self.signature = []
class OC:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
        self.suffix = []
class Employer:
    def __init__(self):
        self.obj_name = []
        self.org_name = []
        self.suffix = []
        self.address = []
        self.tel = []
class Carrier:
    def __init__(self):
        self.obj_name = []
        self.org_name = []
        self.suffix = []
        self.address = []
class ServicingAgent:
    def __init__(self):
        self.obj_name = []
        self.org_name = []
        self.suffix = []
        self.address = []
class Adjuster:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.job_title = []
        self.org_name = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
class EmployerRep:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.job_title = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
class Entity:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.branch = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.entitytype = []
        self.registered_agent = []
class Claimant:
    def __init__(self): 
        self.obj_name = []
        
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.SSN = []
        self.DOB = []
        self.tel = []
        self.email = []
        
        self.defendant = []
        self.employer = []
        self.carrier = []
        self.serv_agent = []
        self.OC = []
        self.lawyer = []
        self.adjuster = []
        self.employer_rep = []
        self.cc = []
        
        self.judge = []
        self.case_number = []
        self.claim_number = []
        self.accident_date = []
        
        self.case_type = []
        self.state = []
        self.court = []

        self.speciallang = []
class SpecialLanguage:
    def __init__(self): 
        self.obj_name = []
        self.language_name = []
        self.content = []
        self.doc_type = []
class RegisteredAgent:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.branch = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
#Classes
   
   
#Functions  
def ExportValueToKey(key, value, doc, doc_path):
    '''Exports a single value to all instances of a single key in a given document.'''

    key_located = False
    for paragraph in doc.paragraphs: #Scan each paragraph for the provided key
        key_located = True
        if key in paragraph.text: #If the key is found in a given paragraph, iterate through the paragraph's runs and replace each instance of it 
            key_total = paragraph.text.count(value)
            key_replaced = False
            while key_replaced == False:
                for run in paragraph.runs:
                    if value in run.text:
                        run.text = run.text.replace(key, value)
                        key_index += 1     
                if key_index != key_total: #If every instance of the key has NOT been replaced, then the keys are likely formatted incorrectly and need to be reformatted to function
                    print("Error: The following key was formatted wrong: " + key + " \nReformatting the key now...")
                    FixKeyFormatting(key, doc, doc_path)
                else: 
                    key_replaced = True      
    if key_located == False: #If the key was NEVER located ONCE in ANY of the document's paragraphs, then it likely does not exist at all
        print("The following key was never located: " + key + "\nThe document has not been modified, please provide a valid key")
    else: 
        print("All instances of the following key: " + key + " have sucessfully been replaced with the following value: " + value)
        doc.save(doc_path)

def CopyPasteRunFormatting(copyrun, pasterun):
        #Copies the formatting from the copy run, and pastes the formatting into the paste run
        pasterun.font.all_caps = copyrun.font.all_caps
        pasterun.font.bold = copyrun.font.bold
        
        if(str(copyrun.font.color.rgb) != "None"):
            pasterun.font.color.rgb = RGBColor.from_string(str(copyrun.font.color.rgb))
        else:
            pasterun.font.color.rgb = copyrun.font.color.rgb
        
        if(str(copyrun.font.color.theme_color) != "None"):
            pasterun.font.color.theme_color = RGBColor.from_string(str(copyrun.font.color.theme_color))
        
        pasterun.font.complex_script = copyrun.font.complex_script
        pasterun.font.cs_bold = copyrun.font.cs_bold
        pasterun.font.cs_italic = copyrun.font.cs_italic
        pasterun.font.double_strike = copyrun.font.double_strike
        pasterun.font.emboss = copyrun.font.emboss
        pasterun.font.hidden = copyrun.font.hidden
        pasterun.font.highlight_color = copyrun.font.highlight_color
        pasterun.font.imprint = copyrun.font.imprint
        pasterun.font.italic = copyrun.font.italic
        pasterun.font.math = copyrun.font.math
        pasterun.font.name = copyrun.font.name
        pasterun.font.no_proof = copyrun.font.no_proof
        pasterun.font.outline = copyrun.font.outline
        pasterun.font.rtl = copyrun.font.rtl
        pasterun.font.size = copyrun.font.size
        pasterun.font.shadow = copyrun.font.shadow
        pasterun.font.small_caps = copyrun.font.small_caps
        pasterun.font.snap_to_grid = copyrun.font.snap_to_grid
        pasterun.font.spec_vanish = copyrun.font.spec_vanish 
        pasterun.font.strike = copyrun.font.strike
        pasterun.font.subscript = copyrun.font.subscript
        pasterun.font.superscript = copyrun.font.superscript
        pasterun.font.underline = copyrun.font.underline
        pasterun.font.web_hidden = copyrun.font.web_hidden
        
def FixKeyFormatting(key, doc, doc_path):
    runstxt = []
    for paragraph in doc.paragraphs:
            if key in paragraph.text:
                
                #collect the data from each run
                for run in paragraph.runs:
                    runstxt.append(str(run.text))

                #find the start of the key and the end of the key, which is fragmented across muiltiple runs
                runID = 0
                foundleftparam = False
                foundrightparam = False
                rightparampos = ["runtxt", "rightparampos", "runID"]
                leftparampos = ["runtxt", "leftparampos", "runID"]
              
                for runtxt in runstxt:
                    if foundleftparam == False:  
                        foundleftparam = runtxt.find(param_start_indicator)
                        if foundleftparam == -1 and foundleftparam:
                            foundleftparam = False

                    if str(foundleftparam) != "False" and str(foundleftparam) != "True":
                       leftparampos = [runtxt, foundleftparam, runID]
                       foundleftparam = True
                       
                    if str(foundleftparam) == "True":
                        foundrightparam = runtxt.find(param_end_indicator)
                        
                        if foundrightparam == -1 and str(foundrightparam) != "True":
                            foundrightparam = False
                  
                        else:
                            rightparampos = [runtxt, foundrightparam, runID]
                            foundrightparam = True
                    runID += 1
                    
                    if str(foundrightparam) == "True" and str(foundleftparam) == "True":
                        break
                
                #If there is more than two fragments, compile all the fragments inbetween into an array
                leftparam_runID = leftparampos[2]
                rightparam_runID = rightparampos[2]
                betweenparampos = [] #[runtxt, run, runID]
                if leftparam_runID + 1 != rightparam_runID:
                    runID = leftparam_runID + 1
                    while runID < rightparam_runID:
                        betweenparampos.append([paragraph.runs[runID].text, paragraph.runs[runID], runID])
                        runID += 1
                    
                #Is there is noise before the left param, and after the right param? 
                leftparam_runlength = len(leftparampos[0])
                rightparam_runlength = len(rightparampos[0])
                
                leftparam_leftmostindex = 0
                rightparam_rightmostindex = rightparam_runlength - 1
                
                leftparam_index = leftparampos[1]
                rightparam_index = rightparampos[1]
                
                leftparam_noise = False
                rightparam_noise = False
                
                if leftparam_runlength > 1:
                    if leftparam_index != leftparam_leftmostindex:
                        leftparam_noise = True
                if rightparam_runlength > 1:
                    if rightparam_index != rightparam_rightmostindex:
                        rightparam_noise = True

                #How many fragments are there?
                twofragments = False
                threefragments = False
                morethan_threefragments = False
                
                if len(betweenparampos) == 0:
                    twofragments = True
                elif len(betweenparampos) == 1:
                    threefragments = True
                elif len(betweenparampos) > 1:
                    morethan_threefragments = True
                    
                #If there are only two fragments , add an extra run to contain the key 
                if twofragments == True:
                    paragraph.add_run("")

                #If there are only two fragments, create a new string of runs to replace the olds runs, so that the formatting can remain identical despite the disrupting addition of a new run 
                key_runID = "Undefined"
                if twofragments == True:
                    key_runID = rightparam_runID 
                    runID = 0
                    for run in paragraph.runs: 
                        if runID == key_runID:
                            paragraph.runs[runID].text = ""
                        elif runID > key_runID: 
                            original_run = paragraph.runs[runID - 1]
                            replacement_run = paragraph.add_run("") 
                            CopyPasteRunFormatting(original_run, replacement_run) 
                            replacement_run.text = original_run.text
                        runID += 1
                    oldruns_endID = runID
                    runID = 0
                    for run in paragraph.runs:
                        if runID < oldruns_endID and runID > key_runID:
                            paragraph.runs[runID].text = ""
                        runID += 1
               #If there are noisy fragments, extract the noise
                leftparam_noise_content = ""
                rightparam_noise_content = ""
                leftparam_content = leftparampos[0]
                rightparam_content = rightparampos[0]
                
                if leftparam_noise == True:
                    leftparam_noise_content = leftparam_content[0:leftparam_index]
                if rightparam_noise == True:
                    rightparam_noise_content = rightparam_content[rightparam_index + 1:rightparam_runlength]
                       
                #If there is noise, put the noise into a seperate run from their associated fragment and clean up the runs accordingly
                leftparam_runID = leftparampos[2]
                rightparam_runID = rightparampos[2]

                if twofragments == True:  
                    leftparam_runID = leftparampos[2]
                    rightparam_runID = rightparampos[2] + 1
                else:
                    leftparam_runID = leftparampos[2]
                    rightparam_runID = rightparampos[2]
                
                runID = 0
                for run in paragraph.runs:
                    if runID == leftparam_runID:
                        run.text = leftparam_noise_content
                    elif runID > leftparam_runID and runID < rightparam_runID:
                        run.text = ""
                    elif runID == rightparam_runID:
                        run.text = rightparam_noise_content
                        if twofragments == True:
                            CopyPasteRunFormatting(paragraph.runs[runID - 1], run) #Fixes the formatting of the noise, since the formatting for this area is tricky to resolve in the previous code intended for this purpose 
                        break 
                    runID += 1
                    
                #Figure out the appropriate run lengths for each fragment for comparison purposes, if there is noise, subtract it
                if leftparam_noise == True:
                    leftparam_runlength = abs(leftparam_runlength - len(leftparam_noise_content))
                if rightparam_noise == True:
                    rightparam_runlength = abs(rightparam_runlength - len(rightparam_noise_content))
               
                #Determine the appropriate location for the key to be placed 
                if twofragments == True:
                    key_runID = key_runID #Was already set to the newly added run in a previous function  
                else:
                    key_runID = betweenparampos[0][2] #Set to the first inbetween param
                    

                #Figure out the most prominent formatting used for the key, and assume that this is the intended formatting
                if twofragments == True:
                    highest_formatting_prominence = max(leftparam_runlength, rightparam_runlength)
                    if highest_formatting_prominence == leftparam_runlength:  
                        CopyPasteRunFormatting(paragraph.runs[leftparam_runID], paragraph.runs[key_runID])
                    elif highest_formatting_prominence == rightparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[rightparam_runID], paragraph.runs[key_runID])
                        
                if threefragments == True: 
                    betweenparam_runlength = len(betweenparampos[0][0])
                    betweenparam_runID = betweenparampos[0][2]
                    
                    highest_formatting_prominence = max(leftparam_runlength, betweenparam_runlength, rightparam_runlength)
                    if highest_formatting_prominence == leftparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[leftparam_runID], paragraph.runs[key_runID])
                    elif highest_formatting_prominence == betweenparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[betweenparam_runID], paragraph.runs[key_runID])
                    elif highest_formatting_prominence == rightparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[rightparam_runID], paragraph.runs[key_runID])

                if morethan_threefragments == True:
                    betweenparam_runlength = [] #[run length]
                    for betweenparam in betweenparampos:
                        betweenparam_runlength.append(len(betweenparam[0]))
                    highest_betweenparam_runlength = max(betweenparam_runlength)
                    for betweenparam in betweenparampos:
                        if len(betweenparam[0]) == highest_betweenparam_runlength:
                            betweenparam_runID = betweenparam[2]
                            betweenparam_runlength = len(betweenparam[0])
                            break          
                    highest_formatting_prominence = max(leftparam_runlength, highest_betweenparam_runlength, rightparam_runlength)
                    
                    if highest_formatting_prominence == leftparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[leftparam_runID], paragraph.runs[key_runID])
                    elif highest_formatting_prominence == betweenparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[betweenparam_runID], paragraph.runs[key_runID])
                        print(paragraph.runs[key_runID].font.color.rgb)
                    elif highest_formatting_prominence == rightparam_runlength:
                        CopyPasteRunFormatting(paragraph.runs[rightparam_runID], paragraph.runs[key_runID])
                        
                #Add the key into its now correctly formatted placeholder
                paragraph.runs[key_runID].text = key
                    
                #Save the documentcontainer_runID
                doc.save(docPATH)
   
def ExportMainToBackup():
    copy(main_script, backup_script)
     
def ImportBackup():
    copy(backup_script, main_script)
    copy(backup_script, update_script)

def ImportUpdate():
    ExportMainToBackup()
    copy(update_script, main_script)
    
def ExportMainToUpdate():
    copy(main_script, update_script)
      

 


  


#SOLID functions
def ToDigit(character_string):
    '''
    Transforms a string of text into a string of digits. Each pair of digits is an arbritary value associated with a given letter or non-letter character. Non-letter characters are manually replaced with their chosen equivelant, while letter characters are automatically replaced by looping through the ascii_letters array. 
    '''
    
    #Replaces all instances of these non-letter characters with their digit equivelant
    character_string = character_string.replace(" ", "00")
    character_string = character_string.replace("-", "53")
    character_string = character_string.replace("--", "54")
    
    #Replaces all instances of a given letter with their digit equivelant 
    digit = 0
    for letter in ascii_letters:
        digit += 1
        if digit < 10:
            character_string = character_string.replace(letter, "0" + str(digit))
        else:
            character_string = character_string.replace(letter, str(digit))    
    digit_string = character_string
    
    #Returns the string of digits
    return digit_string
      
def ToCharacter(digit_string):
    '''
    Transforms a string of digits into a string of text. Each digit pair is transformed into its corresponding letter or non-letter character. Non-letter character digit pairs are manually transformed into their equivelants. Letter character digit pairs are automatically transformed into their equivelants by looping through the ascii letters array. 
    '''
    digit_string = str(digit_string)
    digit_length = len(digit_string) * .5 #The length is halved to represent the amount of two-digit pairs
    digit_length = int(digit_length)
    
    #Creates a two-dimensional array that pairs a given letter with their digit equivelant by looping through the ascii letter array
    letters = ["null"] #Index 0 is declared null so each digit pair, which starts at 01, lines up with its corresponding letter array index 
    digit_equivelant = 0
    for letter in ascii_letters:
        digit_equivelant += 1
        if digit_equivelant < 10:
            letters.append(["0" + str(digit_equivelant), letter])
        else:
            letters.append([digit_equivelant, letter])
    
    #Creates a string of text, wherein each digit has been replaced with its letter or non-letter equivelant 
    txt = ""
    digit_pair_index = 0
    digit_pair_start = 0
    digit_pair_end = 2
    while digit_pair_index < digit_length:
        digit_pair = digit_string[digit_pair_start:digit_pair_end]

        if digit_pair == "00":
            letter = " "
        elif digit_pair == "53":
            letter = "-"
        elif digit_pair == "54":
            letter = "--"
        else:
            digit_pair = int(digit_pair)
            letter = letters[digit_pair][1] 
        txt += letter
       
        digit_pair_index += 1
        digit_pair_start += 2
        digit_pair_end += 2  
    return txt
    
def GenerateID(obj):
    '''
    This functions generates an ID. Each ID consists of two values: the object type and the object name. The object type is a two digit string that precedes the object name. If the object provided is not a valid object type, then the ID is returned as null. 
    '''
    #Get the type tag associated with the object type, if the object type does not exist, then the type tag and the ID is null
    ID = "null"
    type_tag = "null"
    obj_type = GetObjType(obj, "null")
    if obj_type == "DocType":
        type_tag = 0    #elif_end
    elif obj_type == "Doc":
        type_tag = 1    #elif_end
    elif obj_type == "Lawyer":
        type_tag = 2    #elif_end
    elif obj_type == "OC":
        type_tag = 3    #elif_end
    elif obj_type == "Employer":
        type_tag = 4    #elif_end
    elif obj_type == "Carrier":
        type_tag = 5    #elif_end
    elif obj_type == "ServicingAgent":
        type_tag = 6    #elif_end 
    elif obj_type == "Adjuster":
        type_tag = 7    #elif_end 
    elif obj_type == "EmployerRep":
        type_tag = 8    #elif_end 
    elif obj_type == "Entity":
        type_tag = 9    #elif_end 
    elif obj_type == "Claimant":
        type_tag = 10   #elif_end 
    elif obj_type == "SpecialLanguage":
        type_tag = 11   #elif_end 
    elif obj_type == "RegisteredAgent":
        type_tag = 12   #elif_end 
  #Ensures that null is returned when the object type is invalid 
    if type_tag == "null":
        ID = "null"
        return ID
    else:
        ID = str(type_tag) + ToDigit(getattr(obj, "obj_name")[0])
    #Ensures that the type tag is always two digits
    if type_tag < 10:
        ID = "0" + ID
    return ID
    #END    
def ArrayID(ID):
    '''Retrieves each digit pair that comprises a given ID, and then organizes them into a neatly seperated array'''
    
    #Iterates through each digit pair in the ID, and appends the digit pair to the digit pairs array
    ID = str(ID)
    ID_length = len(ID) * .5 #The length is halved to represent the amount of two-digit pairs
    digit_pair_start = 0
    digit_pair_end = 2
    digit_pairs = []
    digit_pair_index = 0
    while digit_pair_index < len(ID):
        digit_pair = ID[digit_pair_start:digit_pair_end]
        digit_pairs.append(digit_pair)
        digit_pair_start += 2
        digit_pair_end += 2
        digit_pair_index += 2
    return digit_pairs

def CreateEmptyObj(obj_type):
    '''Returns an empty object of the desired type'''
    obj = "null"
    if obj_type == "DocType":
        obj = DocType()
    elif obj_type == "Doc":
        obj = Doc()
    elif obj_type == "Lawyer":
        obj = Lawyer()
    elif obj_type == "OC":
        obj = OC()
    elif obj_type == "Employer":
        obj = Employer()
    elif obj_type == "Carrier":
        obj = Carrier()
    elif obj_type == "ServicingAgent":
        obj = ServicingAgent()
    elif obj_type == "Adjuster":
        obj = Adjuster()
    elif obj_type == "EmployerRep":
        obj = EmployerRep()
    elif obj_type == "Entity":
        obj = Entity()
    elif obj_type == "Claimant":
        obj = Claimant()
    elif obj_type == "SpecialLanguage":
        obj = SpecialLanguage()
    elif obj_type == "RegisteredAgent":
        obj = RegisteredAgent()
    return obj
    #END  
def GetObjKeys(ID = "null", obj = "null", obj_type = "null"):
    #Get the object from the object type 
    if obj != "null":
        obj_type = GetObjType(obj)
        obj = CreateObj(obj_type)
    elif ID != "null":
        obj_type = GetObjType(ID)
        obj = CreateObj(obj_type)
    elif obj_type != "null":
        obj = CreateEmptyObj(obj_type)
    

    #Get the keys associated with an object, and store them as a single cleaned up string
    keys = str(obj.__dict__.keys())
    keys = keys.replace("dict_keys", "")
    keys = keys.replace("(", "")
    keys = keys.replace(")", "")
    keys = keys.replace("[", "")
    keys = keys.replace("]", "")
    keys = keys.replace("'", "")
    
    #Deconstruct the string of keys into an array of keys 
    key_start = 0
    key_end = 0
    skip_identifier = len(", ")
    
    key_index = 0
    key = "null"
    key_array = []
    while key_index < keys.count(", ") + 1:
        if key_index != 0:
            key_start = keys.find(", ", key_end)
                        
        if key_start != 0:
            key_end = keys.find(", ", key_start + skip_identifier)
        else:
            key_end = keys.find(", ", key_start)
                        
        if key_end == -1:
            key_end = len(keys)
                        
        if key_index != 0:
            key = keys[key_start + skip_identifier:key_end]
        else:
            key = keys[key_start:key_end]
        key_array.append(key)
        key_index += 1
    return key_array  

def GetKeyColumns(obj_type, obj_type_keys = "null"):
    '''
    Using the amount of keys associated with an object type, return an array with an equal amount of columns. 
    Each key is correlated to a column by index, for example, the first key of an object type is correlated to the first column of the spreadsheet.
    (starting at column B, since the A column is always designated to IDs)
    '''
    
    #Get the keys of a given object type and then get  thecolumns to store them, or get the columns that are already storing them, as both are the same
    if obj_type_keys == "null":
        obj_type_keys = GetObjKeys("null", "null", obj_type)
    key_amount = len(obj_type_keys)
    letters = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q", "R", "S", "T", "U", "V", "W", "X", "Y", "Z"]
    columns = []
    column_array_complete = False
    while column_array_complete == False: 
        for letter in letters:
            if letter != "A":
                columns.append(letter)
            column_amount = len(columns)
            if column_amount == key_amount:
                break       
        if column_amount == key_amount:
            break
            
        for first_letter in letters:
            if column_amount == key_amount:
                break
            for second_letter in letters:
                columns.append(first_letter + second_letter)
                column_amount = len(columns)
                if column_amount == key_amount:
                    break
        if column_amount == key_amount:
            break
        
        last_column = "null" #At this stage, the last column must be checked because the last column in an excel spreadsheet is XFD
        for first_letter in letters:
            if column_amount == key_amount or last_column == "XFD":
                break
            for second_letter in letters:
                if column_amount == key_amount or last_column == "XFD":
                    break
                for third_letter in letters:
                    columns.append(first_letter + second_letter + third_letter)
                    column_amount = len(columns)
                    last_column_index = column_amount - 1
                    last_column = columns[last_column_index]
                    if column_amount == key_amount or last_column == "XFD":
                        break
        if column_amount != key_amount:
            print("Error: There are more than 16,384 keys for this object type, which exceeds the amount of rows available in an excel spreadsheet.\nAny key past the 16,384th key has not been saved.\nPlease shorten the amount of keys associated with this object type for this object type to work as intended.")
        if column_amount == key_amount or last_column == "XFD1":
            break
    return columns
    
def GetObjType(obj = "null", ID = "null"):
    '''
    Get an object type by using either an ID or an object.
    If the ID is provided, identify the object type using the ID's two digit type tag, otherwise simply get the object's name
    '''
    
    if obj != "null":
        obj_type = str(type(obj).__name__)
    elif ID != "null":
        obj_type = "null"
        type_tag = str(ID[0:2])
        obj_type_identified = False
        if type_tag == "00":
            obj_type = "DocType"#elif_end       
        elif type_tag == "01":
            obj_type = "Doc"#elif_end 
        elif type_tag == "02":
            obj_type = "Lawyer"#elif_end
        elif type_tag == "03":
            obj_type = "OC"#elif_end 
        elif type_tag == "04":
            obj_type = "Employer"#elif_end
        elif type_tag == "05":
            obj_type = "Carrier"#elif_end
        elif type_tag == "06":
            obj_type = "ServicingAgent"#elif_end
        elif type_tag == "07":
            obj_type = "Adjuster"#elif_end
        elif type_tag == "08":
            obj_type = "EmployerRep"#elif_end
        elif type_tag == "09":
            obj_type = "Entity"#elif_end
        elif type_tag == "10":
            obj_type = "Claimant"#elif_end
        elif type_tag == "11":
            obj_type = "SpecialLanguage"#elif_end
        elif type_tag == "12":
            obj_type = "RegisteredAgent"#elif_end 
       
    return obj_type
    #END    
def CreateSaveFile(obj_type, obj_type_keys = "null"):
    '''
    This function creates a new save file or overwrites an existing save file for a given object type by: 
    creating a new excel spreadsheet via openpyxl, creating an ID column, and lastly assigning each of the object type's keys a column
    '''
    #Create the excel spreadsheet to store data pertaining to an object's type
    save_file_name = obj_type.lower() + "_data"
    save_file_PATH = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "data\\" + save_file_name + ".xlsx")
    save_file = Workbook()
    save_file.save(save_file_PATH)
    
    #Add the ID column to the spreadsheet 
    input_data = save_file.active
    input_data["A1"] = "ID"
    
    #Get an array of columns that can be used to store the keys
    columns = GetKeyColumns(obj_type, obj_type_keys)
    
    #Assign each key a column
    if obj_type_keys == "null":
        obj_type_keys = GetObjKeys("null", "null", obj_type)
    key_index = 0
    for key in obj_type_keys:
        key_column = columns[key_index] + "1" #All the key columns are in the first row of a given spreadsheet 
        input_data[key_column] = key
        key_index += 1
    save_file.save(save_file_PATH)
    print("created a save file for the " + obj_type + " object type")

def GetSaveFile(obj_type):
    '''This function gets an object type's save file by inputting an object type into the standard path for all save files'''
    save_file_name = obj_type.lower() + "_data"
    save_file = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "data\\" + save_file_name + ".xlsx")
    return save_file 

def GetObj(ID):
    '''
    This function uses an ID to retrieve an object from its associated save file. It does this by:
        First, getting the object's type, getting the object's save file using its type, and then preparing the save file to be read
        Second, iterating through the ID column's rows until the ID is found. This row index is the row index for the entire object
        Third, it iterates through the key columns of the row and retrieves the values associated with each
        Fourth, it creates a new object to assign the values to 
        Fifth, it iterates through each value and organizes its sub-values into an array
        Sixth, it creates a new object and then assigns each value array to its associated key
        Seventh, it returns the workable object 
    '''
    obj = "null"
    
    #Get the object type, get the save file for the object type, and then prepare the save file to be read 
    obj_type = GetObjType("null", ID)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    output_data = save_file.active
    
    #Find the ID's row in the ID column and then use it to retrieve the object's row   
    row_index = 2
    ID_cell = "null"
    ID_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    while ID_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if output_data[cell].value == ID:
            ID_cell_found = True
            ID_cell = cell
            obj_row_index = row_index 
            break 
        row_index += 1
    if ID_cell == "null":
        print("Error: object does not exist or object's ID is corrupted")
        return "null"
    
    #First, iterate through the object's keys and retrieve the values associated with each
    #Second, iterate through the value and organize its sub-values into an array by utilizing the indicators that seperate them 
    #Third, assign each value array to its associated key, and then return a workable object once complete 
    key_index = 0
    obj = CreateEmptyObj(obj_type)
    obj_keys = GetObjKeys("null", "null", obj_type)
    columns = GetKeyColumns(obj_type)
    for key in obj_keys:
        key_column = columns[key_index] + str(obj_row_index)
        value = output_data[key_column].value
        if str(value) != "None":
            total_values = value.count(value_start_indicator) - 1 #Index from zero not from one, do this by subtracting one from the total amount of values
            value_index = 0
            last_value_start = 0
            last_value_end = 0
            skip_start_indicator = len(value_start_indicator)
            skip_end_indicator = len(value_end_indicator)
            value_array = []
            while value_index <= total_values:
                value_start = value.find(value_start_indicator, last_value_start)
                value_end = value.find(value_end_indicator, last_value_end)
                value_data = value[value_start + skip_start_indicator:value_end]
                last_value_start = value_start + skip_start_indicator
                last_value_end = value_end + skip_end_indicator
                value_index += 1
                value_array.append(value_data)
            setattr(obj, key, value_array)
        else:
            setattr(obj, key, "")
        key_index += 1
    print("retrieved object: " + obj.obj_name[0])
    return obj

def SaveObj(obj, new_object = False):
    '''
    This function saves an object to its associated save file by:
        First, getting the object's ID and deleting any pre-existing duplicate 
        Second, getting the object's type, keys, and save file path. 
        Third, iterating through the rows of the ID column until it finds an empty row to store the object. 
        Fourth, storing the object's ID in the ID column 
        Fifth, getting the columns associated with each of the object's keys. 
        Sixth, getting the value array associated with each of the object's keys and then organizing the values into an array format fit for a cell
        Seventh, assigning each key's associated value array to the key's associated column.
        Eighth, saving the spreadsheet.
    '''
    #Generate the object's ID
    ID = GenerateID(obj)
    
    #If the object already exists, delete it. Assume that it might unless otherwise stated 
    if new_object == False:
        DeleteObj(ID)
        
    #Get the object type, get the object type's keys, get the save file for the object type, and prepare the save file to be edited 
    obj_type = GetObjType(obj, "null")
    obj_keys = GetObjKeys("null", "null", obj_type)
    ID = GenerateID(obj)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    input_data = save_file.active
    
    #Find the first empty ID cell to store the object's ID. Its row index will indicate the row where the rest of the object should be stored 
    row_index = 2
    empty_ID_cell = "null"
    empty_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    while empty_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if input_data[cell].value == None:
            input_data[cell].number_format = numbers.FORMAT_TEXT #format the ID cell as text to prevent corruption from undefined formatting 
            empty_cell_found = True
            empty_ID_cell = cell
            empty_row_index = row_index 
            break 
        row_index += 1
    if empty_ID_cell == "null":
        print("Error: There is no avaliable cell that can be used to store this object.")
        return "null"
    
    #Store the ID in the first available empty cell in the A column 
    input_data[empty_ID_cell] = ID
    
    #Store the value array in the column associated with each key 
    #IMPORTANT: This algorithm assumes that the key initiations of a given object type's associated class HAVE NOT CHANGED POSITION since the conception of the save file 
    #IMPORTANT: DO NOT add new key iniations BETWEEN existing iniations, as then the code below will NOT work as intended. ONLY add new iniations to the end of a given class
    columns = GetKeyColumns(obj_type)
    column_index = 0
    for key in obj_keys:
        value_array = getattr(obj, key)
        value_array = str(value_array)
        value_array = value_array.replace("[", "")
        value_array = value_array.replace("]", "")
        value_array = value_array.replace("'", "")
        value_array = value_start_indicator + value_array
        value_array = value_array + value_end_indicator
        value_array = value_array.replace(", ", value_end_indicator + value_start_indicator)
        if value_array != value_start_indicator + value_end_indicator: #If the value array is empty, there is no need to waste storage saving it as empty
            cell = columns[column_index] + str(row_index)
            input_data[cell] = value_array
        column_index += 1
    save_file.save(save_file_PATH)
    print("object saved at row " + str(row_index))

def DeleteObj(ID):
    '''
    This function deletes an object by: 
        First, preparing the save file to be modified 
        Second, using an ID to find the object's row
        Third, delete the object's row using excel's built-in deletion ability 
    '''
    
    #Get the object type, get the object type's keys, get the save file for the object type, and prepare the save file to be edited 
    obj_type = GetObjType("null", ID)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    delete_data = save_file.active 
    
    #Find the ID's row in the ID column and then use it to retrieve the object's row   
    row_index = 2
    ID_cell = "null"
    ID_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    while ID_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if delete_data[cell].value == ID:
            ID_cell_found = True
            ID_cell = cell
            obj_row_index = row_index 
            break 
        row_index += 1
    if ID_cell == "null":
        print("Error: object does not exist or object's ID is corrupted")
        return "null"
    delete_data.delete_rows(row_index, 1)
    save_file.save(save_file_PATH)
    print("deleted object at row " + str(row_index))
    
def GetSimilarIDs(ID, similarity_percentage):
    '''
    Get similar IDs within the specified similarity range. Do this by:
        First, preparing the save file to be read
        Second, fetching an ID from a given ID row
        Third, comparing the digit pairs of this comparand ID to the provided ID
        Fourth, dividing the amount of identified similar digit pairs by the amount of total digit pairs
        Fifth, checking to see if this percentage matches the desired similarity percentage 
        Sixth, adding it to the similar IDs array if it does 
        Seventh, iterating to the next row until there are no IDs left to compare 
        Eigth, returning the similar IDs array 
    '''
    #Get the object type, get the save file for the object type, and then prepare the save file to be read 
    obj_type = GetObjType("null", ID)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    output_data = save_file.active
    
    #Iterate through each row in the ID column until there are no IDs left
    row_index = 2
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    ID_digit_pairs = ArrayID(ID)
    similar_IDs = []
    while row_index != row_max:
        cell = "A" + str(row_index)
        comparand_ID = output_data[cell].value
        if str(comparand_ID) == "None":
            break
        comparand_digit_pairs = ArrayID(comparand_ID) 
        similar_pairs = 0
        #Once a comparand ID has been retrieved compare its digit pairs to the provided IDs digit pairs
        for comparand_digit_pair in comparand_digit_pairs:
            for ID_digit_pair in ID_digit_pairs:
                if comparand_digit_pair == ID_digit_pair:
                    similar_pairs += 1
        similarity = similar_pairs/len(comparand_digit_pairs)
        #If the percentage of similar digit pairs in the comparand ID matches the desired similarity percentage, add it to the similar ID array
        if similarity >= similarity_percentage:
            similar_IDs.append(comparand_ID)
        row_index += 1
    print("found " + str(len(similar_IDs)) + " similar IDs")
    return similar_IDs

def MatchDataToIDs(data, obj_type, data_type = "null"):
    '''
    Use data to find correlated IDs, and if avaliable use the specified data type to speed up the search. Do this by:
        First, preparing the save file to be read.
        Second, if the data type was specified, identify its correlated key column, otherwise iterate though every key column 
        Third, iterate through the rows of a given key column and add rows with a match to an array
        Fourth, iterate through the array of matching rows and add their correlated IDs to an array
        Fifth, return the ID array
    '''
    
    #Get the object type, get the save file for the object type, and then prepare the save file to be read 
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    output_data = save_file.active
    
    data_column = "null"
    keys = GetObjKeys("null", "null", obj_type)
    key_columns = GetKeyColumns(obj_type)
    
    #If the data type was specified, identify which key column correlates to the data type
    if data_type != "null":
        key_index = 0
        for key in keys:
            key = str(key)
            is_data_type = key.find(data_type)
            if is_data_type != -1:
                data_type_index = key_index
                break
            else:
                key_index += 1
        data_column = key_columns[data_type_index]

        #Iterate through the data type's rows and add each matching row to an array
        row_index = 2
        data_match_rows = []
        row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
        while row_index != row_max:
            cell = data_column + str(row_index)
            cell_value = output_data[cell].value
            if str(cell_value) != "None":
                is_data = cell_value.find(data)
                if is_data != -1:
                    data_match_rows.append(row_index) 
            row_index += 1
        
        #Iterate through the array of matching rows and identify the IDs of the objects that contain the matching data
        ID_matches = []
        for row in data_match_rows:
            cell = "A" + str(row)
            ID = output_data[cell].value
            ID_matches.append(ID)
            
    #Iterate through all key columns of an object type if a data type is not specified 
    else:
        data_match_rows = []
        row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
        #Iterate through the rows of every key column and add any matching rows to an array
        for column in key_columns:
            print("searching column: " + column)
            row_index = 2
            while row_index != row_max:
                cell = column + str(row_index)
                cell_value = output_data[cell].value
                if str(cell_value) != "None":
                    is_data = cell_value.find(data)
                    if is_data != -1:
                        data_match_rows.append(row_index) 
                row_index += 1
                
        #Iterate through the array of matching rows and identify the IDs of the objects that contain the matching data
        ID_matches = []
        for row in data_match_rows:
            cell = "A" + str(row)
            ID = output_data[cell].value
            ID_matches.append(ID)
    return ID_matches
        
    #Find the ID's row in the ID column and then use it to retrieve the object's row   
    row_index = 2
    data_cell = "null"
    data_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    '''
    while data_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if output_data[cell].value == ID:
            ID_cell_found = True
            ID_cell = cell
            obj_row_index = row_index 
            break 
        row_index += 1
    if ID_cell == "null":
        print("Error: object does not exist or object's ID is corrupted")
        return "null"
    '''

def CreateObjType(obj_type, obj_type_keys): 
    '''This function creates new objects automatically. Since this function updates existing code, it's work is first saved in the update file. To enact the update it must be imported to the main file via ImportUpdate(). Should you wish to remove the update, you can invoke ImportBackup(), which will revert both the update and main file back to their previous renditions.'''
    
    #Export the latest version of the main script to the extension file 
    ExportMainToUpdate()
    
    single_indent = "    "
    double_indent = "        "
    triple_indent = "            "
    with open(update_script, "r") as update:
        update_txt = update.read() #Get all of the text in the update file as a string
        
        #Create a class declaration for this object and iniate the desired keys 
        classes_start = update_txt.find("#Classes") + len("#Classes")
        classes_end = update_txt.find("#Classes", classes_start)
        classes = update_txt[classes_start:classes_end] #The position where class declarations occur 
        classes = classes + "class " + obj_type + ":" + newline + single_indent + "def __init__(self):" + newline + double_indent #Declaring the class
        classes = classes + "self." + "obj_name" + " = []" + newline + double_indent#Add the object name key to all classes
        for key in obj_type_keys: #Iterating through the desired keys and iniating them 
            if key == obj_type_keys[len(obj_type_keys) - 1]:
                classes = classes + "self." + key + " = []" + newline #prevents the classes comment from being formatted incorrectly 
            else:
                classes = classes + "self." + key + " = []" + newline + double_indent
        update_first_half = update_txt[0:classes_start]
        update_second_half = update_txt[classes_end:len(update_txt)]
        update_txt = update_first_half + classes + update_second_half #Saving the changes 
        
        #Add the object's type tag to the Generate ID function
        GenerateID_start = update_txt.find("def GenerateID(obj):")
        GenerateID_end = update_txt.find("#END", GenerateID_start)
        GenerateID = update_txt[GenerateID_start:GenerateID_end]
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = GenerateID.find("elif obj_type ==", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        last_elif_start = last_elif_pos
        last_type_tag_pos = GenerateID.find("type_tag = ", last_elif_start) + len("type_tag = ")
        last_type_tag = GenerateID[last_type_tag_pos:last_type_tag_pos + 2]#Get the two digit type type
        last_type_tag = int(last_type_tag)
        type_tag = str(last_type_tag + 1)
        end_indicator_length = len("#elif_end")
        newline_length = len(newline)
        last_elif_end = GenerateID.find("#elif_end", last_elif_start) + end_indicator_length + newline_length
        GenerateID = GenerateID[0:last_elif_end] + newline + single_indent + "elif obj_type == " + '"' + obj_type + '"' + ":" + newline + double_indent + "type_tag = " + type_tag + "   #elif_end" + GenerateID[last_elif_end:len(GenerateID)]#Add the new type tag to the elif expression      
        update_first_half = update_txt[0:GenerateID_start]
        update_second_half = update_txt[GenerateID_end:len(update_txt)]
        update_txt = update_first_half + GenerateID + update_second_half#Save the changes made to GenerateID()
        
        #Add the object's type tag to the GetObjType() function 
        GetObjType_start = update_txt.find("def GetObjType(obj = ")
        GetObjType_end = update_txt.find("#END", GetObjType_start)
        GetObjType = update_txt[GetObjType_start:GetObjType_end]
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = GetObjType.find("elif type_tag == ", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        last_elif_start = last_elif_pos
        last_elif_end = GetObjType.find("#elif_end", last_elif_start) + len("#elif_end") + len(newline)
        GetObjType = GetObjType[0:last_elif_end] + newline + double_indent + "elif type_tag == " + '"' + type_tag + '"' + ":" + newline + triple_indent + "obj_type = " + '"' + obj_type + '"' + "#elif_end" + GetObjType[last_elif_end:len(GetObjType)]#Add the new type to the elif expression
        update_first_half = update_txt[0:GetObjType_start]
        update_second_half = update_txt[GetObjType_end:len(update_txt)]
        update_txt = update_first_half + GetObjType + update_second_half#Save the changes made to GetObjType()
        
        #Add the object's type to the CreateEmptyObj() function
        CreateEmptyObj_start = update_txt.find("def CreateEmptyObj(obj_type)")
        CreateEmptyObj_end = update_txt.find("#END", CreateEmptyObj_start)
        CreateEmptyObj = update_txt[CreateEmptyObj_start:CreateEmptyObj_end]
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = CreateEmptyObj.find("elif obj_type == ", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        last_elif_start = last_elif_pos
        last_elif_end = CreateEmptyObj.find("()", last_elif_start) + len("()")
        CreateEmptyObj = CreateEmptyObj[0:last_elif_end] + newline + single_indent + "elif obj_type == " + '"' + obj_type + '"' + ":" + newline + double_indent + "obj = " + obj_type + "()" + CreateEmptyObj[last_elif_end:len(CreateEmptyObj)]
        update_first_half = update_txt[0:CreateEmptyObj_start]
        update_second_half = update_txt[CreateEmptyObj_end:len(update_txt)]
        update_txt = update_first_half + CreateEmptyObj + update_second_half#Save the changes made to GetObjType()
        
        #Create save file
        CreateSaveFile(obj_type, obj_type_keys)
        
        #Overwrite previous update file with the new update 
        with open(update_script, "w") as update: 
            update.write(update_txt)
        print("Created the " + obj_type + " object type. Invoke ImportUpdate() to save its addition.")

def DeleteObjType(obj_type):
    '''
    This function deletes an object type by:
        1. Deleting its class declaration
        2. Deleting it from the CreateEmptyObj() function
        3. Deleting it from the GenerateID() function, and then revising the type tags that follow it to reflect its abscence 
        4. Deleting it from the GetObjType() function, and then revising the type tags that follow it to reflect its abscence
    All updates made to the source code are saved in the update file. Therefore, for the deletion to be enacted, the ImportUpdate() function must be invoked. 
    '''
    
    if obj_type == "Doc":
        print("This object type cannot be deleted, as it is essential for the program to function as intended")
        return
    if CreateEmptyObj(obj_type) == "null":
        print("error: object cannot be deleted, as object does not exist")
        return
        
    ExportMainToUpdate()   
    single_indent = "    "
    double_indent = "        "
    triple_indent = "            "
    with open(update_script, "r") as update:
        update_txt = update.read() #Get all of the text in the update file as a string    
        
        #Delete the class declaration for the object type
        classes_start = update_txt.find("#Classes") + len("#Classes")
        classes_end = update_txt.find("#Classes", classes_start)
        classes = update_txt[classes_start:classes_end] #The position where class declarations occur 
        obj_type_start = classes.find("class " + obj_type)
        obj_keys = GetObjKeys("null", "null", obj_type)
        last_key = obj_keys[len(obj_keys) - 1]
        obj_type_end = classes.find(last_key) + len(last_key) + len(" = []") + len(newline)
        classes = classes[0:obj_type_start] + classes[obj_type_end:len(classes)] #Delete the object type's class declaration by omission 
        update_first_half = update_txt[0:classes_start]
        update_second_half = update_txt[classes_end:len(update_txt)]
        update_txt = update_first_half + classes + update_second_half #Save the deletion 
        
        #Delete the object type from the CreateEmptyObj() function
        CreateEmptyObj_start = update_txt.find("def CreateEmptyObj(obj_type)")
        CreateEmptyObj_end = update_txt.find("#END", CreateEmptyObj_start)
        CreateEmptyObj_ = update_txt[CreateEmptyObj_start:CreateEmptyObj_end]
        obj_type_end = CreateEmptyObj_.find(obj_type + "()") + len(obj_type) + len("()") + len(newline)
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = CreateEmptyObj_[0:obj_type_end].find(single_indent + "elif obj_type == ", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        obj_type_start = last_elif_pos - 1
        CreateEmptyObj_ = CreateEmptyObj_[0:obj_type_start] + CreateEmptyObj_[obj_type_end:len(CreateEmptyObj_)]
        update_first_half = update_txt[0:CreateEmptyObj_start]
        update_second_half = update_txt[CreateEmptyObj_end:len(update_txt)]
        update_txt = update_first_half + CreateEmptyObj_ + update_second_half
        
        #Delete the object type from GenerateID()
        GenerateID_start = update_txt.find("def GenerateID(obj):")
        GenerateID_end = update_txt.find("#END", GenerateID_start)
        GenerateID = update_txt[GenerateID_start:GenerateID_end]
        obj_type_start = GenerateID.find(single_indent + "elif obj_type == " + '"' + obj_type + '"')
        obj_type_end = GenerateID.find("#elif_end", obj_type_start) + len("#elif_end") + len(newline) + len(newline)
        obj_type_pos = GenerateID[obj_type_start:obj_type_end]
        type_tag_start = obj_type_pos.find("type_tag = ") + len("type_tag = ")
        type_tag_end = type_tag_start + 2
        type_tag = obj_type_pos[type_tag_start:type_tag_end]
        type_tag = int(type_tag)
        type_tags_revised = False
        GenerateID = GenerateID[0:obj_type_start] + GenerateID[obj_type_end:len(GenerateID)]
        type_tag_index = 1
        elif_end = GenerateID.find("#Ensures that null is returned when the object type is invalid")
        elif_pos = GenerateID[0:elif_end]
        while type_tags_revised == False: #Iterate through the type tags after the the object type's tag and subtract one from them to reflect the one less type tag
            type_tag_to_revise_start = elif_pos.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 2
            type_tag_to_revise_end = elif_pos.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 3
            type_tag_to_revise = elif_pos[type_tag_to_revise_start] + elif_pos[type_tag_to_revise_end]
            try:
                type_tag_to_revise = int(type_tag_to_revise) #If the integer is invalid then there are no more type tags to revise. Therefore, break the loop
                GenerateID = GenerateID.replace(str(type_tag_to_revise), str(type_tag_to_revise - 1))
                type_tag_index += 1
            except:
                type_tags_revised = True
                break      
        update_first_half = update_txt[0:GenerateID_start]
        update_second_half = update_txt[GenerateID_end:len(update_txt)]
        update_txt = update_first_half + GenerateID + update_second_half

        #Delete the object type from the GetObjType() function
        GetObjType_start = update_txt.find("def GetObjType(obj = ")
        GetObjType_end = update_txt.find("#END", GetObjType_start)
        GetObjType = update_txt[GetObjType_start:GetObjType_end]
        obj_type_start = GetObjType.find(single_indent + "elif type_tag == " + '"' + str(type_tag) + '"') - len(single_indent)
        if obj_type_start == (-1 - len(single_indent)):
            obj_type_start = GetObjType.find(single_indent + "elif type_tag == " + '"' + "0" + str(type_tag) + '"') - len(single_indent)
        obj_type_end = GetObjType.find("obj_type = " + '"' + obj_type + '"' + "#elif_end") + len("obj_type = " + '"' + obj_type + '"' + "#elif_end" + newline) 
        GetObjType = GetObjType[0:obj_type_start] + GetObjType[obj_type_end:len(GetObjType)]   
        type_tags_revised = False
        type_tag_index = 1
        while type_tags_revised == False: #Iterate through the type tags after the the object type's tag and subtract one from them to reflect the one less type tag
            type_tag_to_revise_start = GetObjType.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 2
            type_tag_to_revise_end = GetObjType.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 3
            type_tag_to_revise = GetObjType[type_tag_to_revise_start] + GetObjType[type_tag_to_revise_end]
            try:
                type_tag_to_revise = int(type_tag_to_revise) #If the integer is invalid then there are no more type tags to revise. Therefore, break the loop
                if type_tag_to_revise - 1 < 10: #If the type tag is less than 10, add a leading zero
                    GetObjType = GetObjType.replace(str(type_tag_to_revise), "0" + str(type_tag_to_revise - 1))
                else:
                    GetObjType = GetObjType.replace(str(type_tag_to_revise), str(type_tag_to_revise - 1))
                type_tag_index += 1
            except:
                type_tags_revised = True
                break    
        update_first_half = update_txt[0:GetObjType_start]
        update_second_half = update_txt[GetObjType_end:len(update_txt)]
        update_txt = update_first_half + GetObjType + update_second_half
        
        #Delete save file
        save_file = GetSaveFile(obj_type)
        os.remove(save_file)

        #Overwrite previous update file with the new update 
        with open(update_script, "w") as update: 
            update.write(update_txt)
        print("Deleted the " + obj_type + "object type. Invoke ImportUpdate() to save its deletion.")
     
def ImportDataFromDoc(doc_path): 
    doc = Document(doc_path)    
#SOLID functions

#def AddCharacterToInputValues(), adds a character to the list of characters allowed to be used for value input by adding the character to the TxtToDigit and DigitToNum functions

#def DeleteObjKey

#def GetDocFromDocObj

#def AddObjKey 

#def ExtractDataFromDoc(), extracts data from other subpoenas      
            
        


#TEST CODE GOES HERE!!!

'''Add deletion of excel file to deletion code but otherwise GOOD'''

'''
doc_path = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "templates\\Doc.docx")
doc = Document(doc_path)




workbook = load_workbook("C:\\Development\\0.0 RawPrograms\\DocGen\\data\\lawyer_data.xlsx")
sheet = workbook.active



sheet["A2"] = "Hello"
sheet["A3"] = "shit"
sheet["B1"] = "world!"

workbook.save("C:\\Development\\0.0 RawPrograms\\DocGen\\data\\lawyer_data.xlsx")

print(sheet["A2"].value)

'''


#TEST CODE GOES HERE!!!

   
#Notes
'''
All Documents have to be a docx document
'''
#Notes
   
   
   
   
   

"""
def SaveObj(obj):
    '''
    This function saves a new object or an updated existing object to its corresponding object type folder. 
    In said folder, objects are added to the object file and the ID file. 
    Before being saved, an ID is generated.
    
    Each ID is composed of the following elements: 
        (a)the name of the object
        (b)the object type
        (c)id position in the id file
        (d)starting object position in the object file
        (e)ending object position in the object file
    
    If the object already exists, the existing object is deleted and replaced with a new object. 
    '''
    ID = "null" 
    
    #Add the object name to the object ID 
    ID = "a"
    
    #Add the object type identifier to the ID
    
    ID = ID + "b"
    obj_type = str(type(obj).__name__)
    match obj_type:
        case "Entity":
            ID = ID + "99"
        case "Lawyer":
            ID = ID + "98" 
        case "Claimant":
            ID = ID + "97" 
        case "OC":
            ID = ID + "96" 
        case "ServAgent":
            ID = ID + "95" 
        case "Adjuster":
            ID = ID + "94" 
        case "SpecialLang":
            ID = ID + "93" 
        case "SubLang":
            ID = ID + "92" 
        case "Doc":
            ID = ID + "91" 
        case "Carrier":
            ID = ID + "90" 
        case "EmployerRep":
            ID = ID + "89" 
        case "Employer":
            ID = ID + "88" 
        case "Misc":
            ID = ID + "87" 
        case "RegisteredAgent":
            ID = ID + "86"
    
    #Insert starting and ending parameter indicators into the object's values so each parameter type and the data associated with it can be distinquished from one another once stored inside the text file
    key_array = GetDictKeys(obj)
    for key in key_array:
        value_array = getattr(obj, key)
        value_array.insert(0, param_start_indicator + key + param_end_indicator)
        value_array.append(param_start_indicator + key + param_end_indicator)
        setattr(obj, key, value_array)    
        
    #Retrieve the save file's associated with the object's type
    save_files = GetObjTypeSaveFiles("null", obj_type)    
    obj_file_PATH = save_files[0]
    ID_file_PATH = save_files[1]
    
    #Find the position where the new ID will be saved in the ID file 
    with open(ID_file_PATH, "r") as ID_file:
        ID_file_PATH = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, str(ID_file.name))
        ID_file_end = os.path.getsize(ID_file_PATH) 
        new_ID_position = str(ID_file_end + 1)

    ID = ID + "c" + new_ID_position
    
    #Find the index where the new object will begin in the object file
    with open(obj_file_PATH, "r") as obj_file:
        obj_file_PATH = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, str(obj_file.name))
        obj_file_end = os.path.getsize(obj_file_PATH) 
        
        new_obj_start = str(obj_file_end + 1)
        ID = ID + "d" + new_obj_start 

    
    #Find the index where the new object will end in the object file. Do this by calculating the amount of bytes the new object will consume. Then, add this total to the amount of bytes currently stored in the object file. This number represents the new index for the last byte in the object file, and thus the ending index for the new object
        #Calculate the amount of bytes the new object will require by determining the amount of characters the new object has, as each character is a single byte
        obj_byte_size = []
        for key in key_array:
            obj_key_values = []
            value_array = getattr(obj, key)
            for value in value_array:
                obj_key_values.append(value)
            obj_byte_size.append(obj_key_values)
        obj_byte_size = str(obj_byte_size)
        obj_byte_size = obj_byte_size.replace("'", "")
        obj_byte_size = obj_byte_size.replace("], [", newline)
        obj_byte_size = obj_byte_size.replace("]", "")
        obj_byte_size = obj_byte_size.replace("[", "")
        obj_byte_size = obj_byte_size.replace("}, {", "}{")
        obj_byte_size = obj_byte_size.replace("}, ", "}")
        obj_byte_size = obj_byte_size.replace(", {", "{")
        obj_byte_size = obj_byte_size + newline + newline + newline + ID + ID + "d"
        obj_byte_size = len(obj_byte_size) + 22 #22 is added to the byte size as the byte size was consistently 22 bytes less than it what should be. Since it is consistently off by this amount, there is no need to diagnose the true cause of this issue, as adding it to the total is much easier. 
        
        #Calculate the amount of characters/bytes the digits of this index will consume once added to the ID by getting the character amount of the byte size variable. If the byte size variable ends in 9, add an extra digit to the total amount of digits, as the byte size variables's character amount will increase by one after the addition of the character amount to the byte size(as 9 + 1 = 10)
        if str(obj_byte_size)[len(str(obj_byte_size)) - 1] != "9":
            obj_byte_size = obj_byte_size + len(str(obj_byte_size))
        else:
            obj_byte_size = obj_byte_size + len(str(obj_byte_size)) 
        
        #Add the object's ending index to the ID. Do this by adding the object's byte size to the previous file size of object file, and then add this total to the ID
        new_obj_end = str(obj_byte_size + obj_file_end) 
        ID = ID + "e" + new_obj_end
        
    with open(obj_file_PATH, "a") as obj_file:
        obj_file.write(newline)
        obj_file.write(ID)
        for key in key_array:
            value_array = getattr(obj, key)
            obj_file.write(newline)
            value_array_index = 0
            for value in value_array:
                obj_file.write(value)
                if value_array_index != 0 and value_array_index != (len(value_array) - 1) and value_array_index != (len(value_array) - 2):
                    obj_file.write(", ")
                value_array_index += 1
        obj_file.write(newline)
        obj_file.write(ID)
        obj_file.write(newline)
     
    obj_file_end = os.path.getsize(obj_file_PATH) 
    print(obj_file_end)
    print(ID)
        
"""
   

DeleteObjType("Pancake")




#Code
'''
dpg.create_context()
dpg.create_viewport(title='DocGen', width=600, height=600)
   
with dpg.font_registry():
    default_font = dpg.add_font("times.ttf", 20)
  
dpg.bind_font(default_font)


with dpg.window(tag="Primary Window"):
    dpg.add_text("Hello, world", pos=(200, 200))
    
   
    
    lawyer =   dpg.add_combo(items=("poo", "poop"), pos=(0,50))
    claimant = dpg.add_combo(items=("poo", "poop"), pos=(0,100))
    entity =   dpg.add_combo(items=("poo", "poop"), pos=(0,150))
    

dpg.setup_dearpygui()
dpg.show_viewport()
dpg.maximize_viewport()

dpg.set_primary_window("Primary Window", True)

dpg.start_dearpygui()
dpg.destroy_context()
'''