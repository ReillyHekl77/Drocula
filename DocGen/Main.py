import dearpygui.dearpygui as dpg
import dearpygui._dearpygui as internal_dpg

import os 
import shutil

from shutil import copy

import string 
from string import ascii_letters

from docx import Document
from docx.shared import RGBColor

from openpyxl import Workbook, load_workbook
from openpyxl.styles import numbers

import time
#Misc Variables
newline = "\n"
tab = "\t"

key_start_indicator = "{"
key_end_indicator = "}"
value_start_indicator = "{"
value_end_indicator = "}"

MAIN_FILE_PATH = os.path.abspath(__file__)
MAIN_FILE_NAME = str(os.path.basename(__file__)) 

backup_script = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "\\backup\\backup.py")
update_script = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "\\update\\update.py")
main_script = MAIN_FILE_PATH
#Misc Variables


#Classes
class MuiltiDoc:
    def __init__(self):
        self.obj_name = []
        self.internal_docs_paths = []
        self.external_docs_paths = []
class Doc:
    def __init__(self):
        self.obj_name = []
        self.doc_path = []     
class Lawyer:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
        self.suffix = []
        self.bar_number = []
        self.signature = []
class OC:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
        self.suffix = []
class Employer:
    def __init__(self):
        self.obj_name = []
        self.org_name = []
        self.suffix = []
        self.address = []
        self.tel = []
class Carrier:
    def __init__(self):
        self.obj_name = []
        self.org_name = []
        self.suffix = []
        self.address = []
class ServicingAgent:
    def __init__(self):
        self.obj_name = []
        self.org_name = []
        self.suffix = []
        self.address = []
class Adjuster:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.job_title = []
        self.org_name = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
class EmployerRep:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.job_title = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.prefix = []
class Entity:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.branch = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
        self.entity_type = []
        self.registered_agent = []
class Claimant:
    def __init__(self): 
        self.obj_name = []
        
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.SSN = []
        self.DOB = []
        self.tel = []
        self.email = []
        
        self.defendant = []
        self.employer = []
        self.carrier = []
        self.serv_agent = []
        self.OC = []
        self.lawyer = []
        self.adjuster = []
        self.employer_rep = []
        self.cc = []
        
        self.judge = []
        self.case_number = []
        self.claim_number = []
        self.accident_date = []
        
        self.case_type = []
        self.state = []
        self.court = []

        self.speciallang = []
class SpecialLanguage:
    def __init__(self): 
        self.obj_name = []
        self.language_name = []
        self.content = []
        self.doc_type = []
class RegisteredAgent:
    def __init__(self):
        self.obj_name = []
        self.first_name = []
        self.middle_name = []
        self.last_name = []
        self.org_name = []
        self.branch = []
        self.address = []
        self.tel = []
        self.fax = []
        self.email = []
#Classes
   


#functions
def CopyPasteRunFormatting(pasterun, copyrun):
        #Copies the formatting from the copy run, and pastes the formatting into the paste run
        pasterun.font.all_caps = copyrun.font.all_caps
        pasterun.font.bold = copyrun.font.bold
        
        if(str(copyrun.font.color.rgb) != "None"):
            pasterun.font.color.rgb = RGBColor.from_string(str(copyrun.font.color.rgb))
        else:
            pasterun.font.color.rgb = copyrun.font.color.rgb
        
        if(str(copyrun.font.color.theme_color) != "None"):
            pasterun.font.color.theme_color = RGBColor.from_string(str(copyrun.font.color.theme_color))
        
        pasterun.font.complex_script = copyrun.font.complex_script
        pasterun.font.cs_bold = copyrun.font.cs_bold
        pasterun.font.cs_italic = copyrun.font.cs_italic
        pasterun.font.double_strike = copyrun.font.double_strike
        pasterun.font.emboss = copyrun.font.emboss
        pasterun.font.hidden = copyrun.font.hidden
        pasterun.font.highlight_color = copyrun.font.highlight_color
        pasterun.font.imprint = copyrun.font.imprint
        pasterun.font.italic = copyrun.font.italic
        pasterun.font.math = copyrun.font.math
        pasterun.font.name = copyrun.font.name
        pasterun.font.no_proof = copyrun.font.no_proof
        pasterun.font.outline = copyrun.font.outline
        pasterun.font.rtl = copyrun.font.rtl
        pasterun.font.size = copyrun.font.size
        pasterun.font.shadow = copyrun.font.shadow
        pasterun.font.small_caps = copyrun.font.small_caps
        pasterun.font.snap_to_grid = copyrun.font.snap_to_grid
        pasterun.font.spec_vanish = copyrun.font.spec_vanish 
        pasterun.font.strike = copyrun.font.strike
        pasterun.font.subscript = copyrun.font.subscript
        pasterun.font.superscript = copyrun.font.superscript
        pasterun.font.underline = copyrun.font.underline
        pasterun.font.web_hidden = copyrun.font.web_hidden

def FormatKeys(doc_path):
    '''
    Since runs can seperate from one another for a plethora of difficult to avoid reasons,
    this function formats the keys as they are intended to be formatted. It does this by:
        1. Finding every paragraph that contains a key, using the prescence of the special key indicators to do so
        2. Iterating through the runs in these key paragraphs and correcting the formatting of each run that requires it
        3. It corrects formatting by seperating any noise from the key
        4. Furthermore, it ensures that the key has a cohesive formatting(the most dominant formatting inputted from the user)
    This function has small limitations that should be stated:
        1. In regard to the key, if there is two equally dominant formats present within it, this function will pick one over the other, despite their equality
        2. If a key is split between paragraphs as opposed to between runs, then this function will not work, the user will have to fix this themselves
            A. This is unlikely, and should be easily avoidable, hence why a solution isn't coded into this function
    '''
    
    #get the document using the provided doc path
    doc = Document(doc_path)
    
    #Find every paragraph with a key in it via the indicators
    key_paragraphs = []
    paragraph_index = 0
    for paragraph in doc.paragraphs:
        if key_start_indicator in paragraph.text and key_end_indicator in paragraph.text:
            key_paragraphs.append(paragraph_index)
        paragraph_index += 1
   
    #Iterate through the paragraphs with keys and ensure their keys are formatted correctly
    for paragraph_index in key_paragraphs:
        #Get the text of every run in the paragraph
        runs_txt = []
        for run in doc.paragraphs[paragraph_index].runs:
            runs_txt.append(run.text)
            
        #Iterate through the text of every run, and add runs with key indicators to an array
        key_runs = []
        run_ID = 0
        for run_txt in runs_txt:
            if value_start_indicator in run_txt:
                indicator_pos = run_txt.find(value_start_indicator)
                key_runs.append([value_start_indicator, run_ID, indicator_pos])
            if value_end_indicator in run_txt:
                indicator_pos = run_txt.find(value_end_indicator)
                key_runs.append([value_end_indicator, run_ID, indicator_pos])
            run_ID += 1

        #Iterate through every key run and flag when a single run contains both the key start and key end indicators 
        flagged_runs = []
        key_run_index = 0
        for key_run in key_runs:
            comparand_key_run_index = 0
            for comparand_key_run in key_runs:
                if key_run[1] == comparand_key_run[1]: #Flag when both indicators COULD exist in the same run 
                    if key_run_index != comparand_key_run_index: #Only flag if the key run and the comparand run are the locations of DIFFERENT indicators 
                        flagged_runs.append(key_run_index)
                comparand_key_run_index += 1
            key_run_index += 1

        #Remove flagged runs from the key runs array, as these runs do not need their formatting changed
        for flagged_run in flagged_runs:
            key_runs[flagged_run] = "flagged"
        flags_removed = False
        while flags_removed == False:
            try:
                key_runs.remove("flagged")
            except ValueError: #If there are no more flags to remove, then break the loop
                flags_removed = True
                break
              
        #If the entire key run array was flagged then skip this paragraph and proceed to the next
        if len(key_runs) < 1:
            paragraph_index += 1
            continue #continue skips this iteration and proceeds to the next
        
        formatting_iteration = 0
        maximum_iteration = len(key_runs)/2
        left_indicator_index = 0
        right_indicator_index = 1
        indicator_pairs = []
        while formatting_iteration < maximum_iteration:
            indicator_pairs.append([key_runs[left_indicator_index], key_runs[right_indicator_index]])
            left_indicator_index += 2
            right_indicator_index += 2
            formatting_iteration += 1
        
        for indicator_pair in indicator_pairs:
            #Assign each member of the indicator pair to its own variable
            left_indicator = indicator_pair[0]
            right_indicator = indicator_pair[1]

            #Assign subvalues to their own variables for readability sake
            left_run_ID = left_indicator[1]
            right_run_ID = right_indicator[1]
            left_indicator_position = left_indicator[2]
            right_indicator_position = right_indicator[2]
            
            #Get the txt of the left indicator's run and the right indicator's run
            left_run_txt = doc.paragraphs[paragraph_index].runs[left_run_ID].text
            right_run_txt = doc.paragraphs[paragraph_index].runs[right_run_ID].text
            
            #Determine whether or not the indicator is enveloped by noise
            left_indicator_noise = False
            if len(left_run_txt) > len(key_start_indicator):
                left_indicator_noise = True
            
            right_indicator_noise = False
            if len(right_run_txt) > len(key_end_indicator):
                right_indicator_noise = True
            
            both_indicator_noise = False
            if left_indicator_noise == True and right_indicator_noise == True:
                both_indicator_noise = True
                
            #Determine the amount of runs between the indicators 
            run_ID_difference = right_run_ID - left_run_ID
            runs_inbetween_indicators = 0
            if run_ID_difference == 1: #There is no run inbetween, the indicators are beside one another
                runs_inbetween_indicators = 0
            elif run_ID_difference == 2: #There is a singular run inbetween the indicators
                runs_inbetween_indicators = 1 
            elif run_ID_difference > 2: #There is an indeterminate amount of runs between the indicators 
                runs_inbetween_indicators = run_ID_difference - 1
            elif run_ID_difference == 0: #Undesired behavior 
                print("ERROR: the format key function is not behaving as intended:" + newline + "It is attempting to correct the formatting of a correctly formatted key.")
                return "null"

            #Create a new run if there is no inbetween run and there is noise present in both indicators
            if runs_inbetween_indicators == 0 and both_indicator_noise == True:
                doc.paragraphs[paragraph_index].add_run("")   
                #Shift the position of each run on the leftside of the right indicator, and the right indicator itself, such that the empty run is placed between the indicators
                run_index = len(doc.paragraphs[paragraph_index].runs) - 1 #the run index starts at the last index in the run array
                key_run_index = right_run_ID
                while run_index >= key_run_index:
                    run = doc.paragraphs[paragraph_index].runs[run_index] #This run
                    copy_paste_run = doc.paragraphs[paragraph_index].runs[run_index - 1] #The run that precedes it
                    run.text = copy_paste_run.text #copy the preceding run's text and paste it into this run
                    CopyPasteRunFormatting(run, copy_paste_run) #copy the preceding run's formatting and paste it into this run
                    run_index -= 1 #rinse and repeat until the runs on the leftside of the right indicator, and the right indicator itself, have all been shifted to the left by one 
                doc.paragraphs[paragraph_index].runs[key_run_index].text = "" #empty the inbetween run of content
                right_run_ID = right_run_ID + 1 #Shift the right run ID to the left by one, as this is where the right indicator now resides 
                runs_inbetween_indicators = 1 #indicate that there is now a singular inbetween run 
            
            #Assign an inbetween run to be the key run
            key_run_ID = right_run_ID - 1
            
            #If there is noise in either indicator, extract the noise and store it in a seperate variable
            left_noise = "null"
            right_noise = "null"
            if left_indicator_noise == True:
                left_noise = left_run_txt[0:left_indicator_position]
            if right_indicator_noise == True:
                right_noise = right_run_txt[right_indicator_position + 1:len(right_run_txt)]

            #Retrieve the key and store its txt in a seperate variable 
            keys_run_txt = "null"
            inbetween_run_ID = "null"
            if runs_inbetween_indicators == 1:
                inbetween_run_ID = left_run_ID + 1
                inbetween_run_txt = doc.paragraphs[paragraph_index].runs[inbetween_run_ID].text
                key_run_txt = left_run_txt[left_indicator_position + 1:len(left_run_txt)] + inbetween_run_txt + right_run_txt[0:right_indicator_position]
            elif runs_inbetween_indicators > 1:
                key_run_txt = left_run_txt[left_indicator_position + 1:len(left_run_txt)]
                inbetween_run_ID = left_run_ID + 1
                while inbetween_run_ID < right_run_ID:
                    inbetween_run_txt = doc.paragraphs[paragraph_index].runs[inbetween_run_ID].text
                    key_run_txt = key_run_txt + inbetween_run_txt
                    inbetween_run_ID += 1
                key_run_txt = key_run_txt + right_run_txt[0:right_indicator_position]
            key_run_txt = key_start_indicator + key_run_txt + key_end_indicator
            
            #Determine the length of each inbetween run, for comparison purposes to see which one's formatting should be dominant 
            inbetween_runs = [] #[run_ID, length]
            if runs_inbetween_indicators == 1: #iterate through the inbetween runs, get their length and run ID
                inbetween_run_ID = left_run_ID + 1
                inbetween_run_txt = doc.paragraphs[paragraph_index].runs[inbetween_run_ID].text
                inbetween_run_length = len(inbetween_run_txt)
                inbetween_runs.append([inbetween_run_ID, inbetween_run_length])
            elif runs_inbetween_indicators > 1:
                inbetween_run_ID = left_run_ID + 1
                while inbetween_run_ID < right_run_ID:
                    inbetween_run_txt = doc.paragraphs[paragraph_index].runs[inbetween_run_ID].text
                    inbetween_run_length = len(inbetween_run_txt)
                    inbetween_runs.append([inbetween_run_ID, inbetween_run_length])
                    inbetween_run_ID += 1
            inbetween_runs_lengths = []
            for inbetween_run in inbetween_runs: #compare the lengths of the inbetween runs to see which one is most dominant
                inbetween_runs_lengths.append(inbetween_run[1])
            dominant_length = max(inbetween_runs_lengths)
            dominant_inbetween_run_ID = "null"
            for inbetween_run in inbetween_runs: #assign the most dominant run to its own variable
                if inbetween_run[1] == dominant_length:
                    dominant_inbetween_run_ID = inbetween_run[0]
                    
            #Determine whether the left indicator's run or the right indicator's run is more dominant 
            left_indicator_run_length = len(left_run_txt[left_indicator_position + 1:len(left_run_txt)]) + len(key_start_indicator)
            right_indicator_run_length = len(right_run_txt[0:right_indicator_position]) + len(key_start_indicator)
            dominant_length = max(left_indicator_run_length, right_indicator_run_length)
            dominant_indicator_run_ID = "null"
            if left_indicator_run_length == dominant_length:
                dominant_indicator_run_ID = left_run_ID
            if right_indicator_run_length == dominant_length:
                dominant_indicator_run_ID = right_run_ID
            
            #Compare the dominant indicator run to the dominant inbetween run
            inbetween_run_length = len(doc.paragraphs[paragraph_index].runs[dominant_inbetween_run_ID].text)
            if left_run_ID == dominant_indicator_run_ID:
                indicator_run_length = left_indicator_run_length
            if right_run_ID == dominant_indicator_run_ID:
                indicator_run_length = right_indicator_run_length    
            dominant_run_length = max(inbetween_run_length, indicator_run_length)
            dominant_run_ID = "null"
            if dominant_run_length == indicator_run_length:
                dominant_run_ID = dominant_indicator_run_ID
                dominant_run = doc.paragraphs[paragraph_index].runs[dominant_run_ID]
            if dominant_run_length == inbetween_run_length:
                dominant_run_ID = dominant_inbetween_run_ID
                dominant_run = doc.paragraphs[paragraph_index].runs[dominant_run_ID]

            #Copy and paste the dominant run's formatting into the key run
            key_run = doc.paragraphs[paragraph_index].runs[key_run_ID]
            CopyPasteRunFormatting(key_run, dominant_run)
            
            #Seperate any noise from its corresponding key by taking the key out of its run
            if left_indicator_noise == True:
                doc.paragraphs[paragraph_index].runs[left_run_ID].text = left_noise
            if right_indicator_noise == True:
                doc.paragraphs[paragraph_index].runs[right_run_ID].text = right_noise
            
            #Clear out any remaining inbetween runs that have not been overwritten
            inbetween_run_ID = left_run_ID + 1
            while inbetween_run_ID != key_run_ID:
                doc.paragraphs[paragraph_index].runs[inbetween_run_ID].text = ""
                inbetween_run_ID += 1
            
            #Clear out any remaining indicator runs that have not been overwritten
            if left_indicator_noise == False:
                doc.paragraphs[paragraph_index].runs[left_run_ID].text = ""
            if right_indicator_noise == False:
                doc.paragraphs[paragraph_index].runs[right_run_ID].text = ""
            
            #Insert the key's text into the empty key run
            key_run.text = key_run_txt

        #Iterate to the next paragraph
        paragraph_index += 1 
    
    #Save the formatting changes
    doc.save(doc_path)
    print("document formatted")

def ExportValueToKey(key, obj, doc_path):
    '''
    This function exports a value to a key by:
        1. Finding the specified key
        2. Determining which value array needs to be retrieved via the specified attribute 
        3. Determining which index, or indices of the array need to be export via the index specifier
        4. Determining if there is any conditional formatting that needs to be evaluated 
        5. Fufilling the instructions of the specifiers and placeholders
    Breakdown of Place Holders and Specifiers:
        1. standard place holders replace all non-character and reserved character values, such as brackets, tabs, newlines, etc.
        2. list place holders iterate through an array, assigning each member of that array to a corresponding value
        3. conditional specifiers check if a given condition is met, if it is, the instruction is executed, otherwise it is voided
    Examples of Key Structure:
        Index Specifier:
            1. {Lawyer.first_name[#~]}
                a. This will iterate through all values in the attribute's value array, as ~ indicates infinity_used_first
            2. {Lawyer.first_name[#-~]}
                b. This will iterate through all values in the attribute's value array, backwards, as the - preceding the ~, indicates reverse
            Note: The index specifier must ALWAYS follow the key's root
                a. {Lawyer.first_name[#~]} -- GOOD
                b. {[#~]Lawyer.first_name} -- BAD
            3. {Lawyer.first_name[#2:4]
                a. iterates through the value array starting from the second value, ending at the fourth value
            4. {Lawyer.first_name[#4:2]
                a. iterates through the value array starting from the fourth value, ending at the second value(back to front)
            5. {Lawyer.first_name[#4]}
                a. only exports the fourth value of the value array
            6. {Lawyer.first_name}
                a. when no index specifier is provided, the value exported defaults to the first one in the value array
            7. {Lawyer.first_name[#-~:10}
                a. iterates through the value array starting from the first value, ending at the tenth value
            8. {Lawyer.first_name[#-~:~}
                a. iterates through the value array starting from the first value, ending at the last value
        Conditional Specifier:
            1. {[if(first)[number_index];]Lawyer.first_name[#~]}
                a. ONLY the first index of the value array is numbered
            2. {[if(-first)[number_index];]Lawyer.first_name[#~]}   
                a. every index BUT the first index of the value array is numbered, as - is the reverse indicator, translates to IF NOT
            3. {[if(-last)[number_index];]Lawyer.first_name[#~]}   
                a. every index BUT the last index of the value array is numbered, as - is the reverse indicator, translates to IF NOT
            4. {[if(last)[number_index];]Lawyer.first_name[#~]}
                a. ONLY the last index of the value array is numbered
            5. {[if(#2)[number_index];]Lawyer.first_name[#~]}
                a. ONLY the second index of the value array is numbered
            6. {[if(#2:6)[number_index];]Lawyer.first_name[#~]}
                a. ONLY the second index through the sixth index of the value array is numbered
            7. {[if(-#2:3)[number_index];]Lawyer.first_name[#~]}
                a. ONLY if the index is NOT 2 or 3, will it be numbered 
            8. {[if(-#2)[number_index];]Lawyer.first_name[#~]}
                a. ONLY if the index is NOT 2, will it be numbered 
            9. {[if(-#~)["This is a note"];]Lawyer.first_name[#~]}
                a. NO index will receive the condition, could be used for including notes that are intended to be deleted upon doc generation
            10. {[if(#~)[number_index];]Lawyer.first_name[#~]}
                a. Every index receives the instruction, this is useless and redudant though included for logic sake
            11. {[if(#-~:~)[number_index];]Lawyer.first_name[#~]}
                a. Every index receives the instruction
            12. {[if(#-~:3)[number_index];]Lawyer.first_name[#~]}
                a. Every index receives the instruction until the third index
            13. [if(#0:3)[number_index];]Lawyer.first_name[#~]}
                a. Every index receives the instruction until the third index
            14. [if(#3:~)[number_index];]Lawyer.first_name[#~]}
                a. Every index receives the instruction after the third index
        Special Character Placeholders:
            These placeholders should be used in place of their counterparts either because,
            their counterparts are non-characters or reserved characters that may otherwise interfere
            with the program's intended behavior. These placeholders are:
                1. [newline]
                2. [tab]
                3. [left_paran]
                4. [right_paran]
                5. [right_bracket]
                6. [left_bracket]
                7. [double_quote]
            Structure:
                1. {[newline]Lawyer.first_name[#~]}
                2. {Lawyer.first_name[#~][newline]}
        Text Placeholders:
            These placeholders will copy paste the specified text to every value.
            {["AKA"][" "]Lawyer.first_name[#~][newline]}
                AKA value_1
                AKA value_2
                AKA value_3
        List Specifiers:
            Structure:
                1. {[xyz_index]Lawyer.first_name[#~]}
                2. {Lawyer.first_name[#~][xyz_index]}
            
            Specifiers:
                1. [xyz_index] -> loops through x, y, z 
                2. [alphabet_index] loops through a, b, c, etc.
                3. [number_index] loops through 1, 2, 3, etc.
                4. [ALPABET_index] loops through A, B, C, etc.
            
            Each of these indexes can be reversed to iterate back to front. For example:
               [-xyz_index] -> loops through z, y, x
            
            Each value of the value array will receive a member of the desired index:
                avalue_1
                bvalue_2
                cvalue_3
            If the user intends there to be more formatting than what is above, they must specify as such
            
            Note, muiltiple list place specifiers can be used:
                {[number_index][alphabet_index][". "]Lawyer.first_name[#~]}
                1a. value_1
                1b. value_2
                1c. value_3    
            
            Lists of the same type being used will also work as follows:
                {[alphabet_index][alphabet_index][". "]Lawyer.first_name[#~]}
                aa. value_1
                bb. value_2
                cc. value_3
        Custom List Specifiers:
            Custom lists can also be specified with the following format:
            {[["i", "ii", "iii"]][". "]Lawyer.first_name[#~]}
            i. value_1
            ii. value_2
            iii. value_3
            These lists behave exactly the same as the standard lists.
            
            They can be iterated through back to front:
            [-["i", "ii", "iii"]][". "]Lawyer.first_name[#~]}
            iii. value_1
            ii. value_2
            i. value_3
            
            They can be used more than once:
            {[["i", "ii", "iii"]][["i", "ii", "iii"]][". "]Lawyer.first_name[#~]}
            ii. value_1
            iiii. value_2
            iiiiii. value_3
         
    ''' 
    place_holders = [["newline", "\n"], ["tab", "\t"], ["double_quote", '"'], ["left_paran", "("], ["right_paran", ")"], ["right_bracket", "]"], ["left_bracket", "["]]
    list_place_holders = ["xyz_index", "alphabet_index", "number_index", "-xyz_index", "-alphabet_index", "-number_index", "ALPHABET_index", "-ALPHABET_index"]
    conditional_specifiers = ["if(first)", "if(last)", "if(-first)", "if(-last)"]
    specifier_start = "["
    specifier_end = "]"
    
    #get the document using the provided doc path
    FormatKeys(doc_path)
    doc = Document(doc_path)
    obj_type = GetObjType(obj) 
    export_key = key #the key structure for a key sending a value
    import_key = obj_type + "." + key #the key structure for a key receiving a value
    
    
    for paragraph in doc.paragraphs:
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:  
                    #if there is an index specifier, replace the key with the desired value or values but ensure that the user's specified formatting is applied
                    if "[#" in run.text: #Is there an index specifier?
                        index_specifier_start = run.text.find("[#")
                        index_specifier_end = run.text.find("]", index_specifier_start) + 1
                        index_specifier = run.text[index_specifier_start:index_specifier_end]
                        export_value_array = getattr(obj, export_key) #export the value array that was saved to the object 
                        if ":" in index_specifier: #The index specifier is a range
                            #Get the first index in the range
                            first_index_start = run.text.find("#") + 1
                            first_index_end = run.text.find(":", first_index_start)
                            first_index = run.text[first_index_start:first_index_end]
                            OG_first_index = "null"
                            if first_index == "~":
                                OG_first_index = first_index
                                first_index = len(export_value_array)
                            if first_index == "-~":
                                OG_first_index = first_index
                                first_index = 1
                            first_index = int(first_index)
                            first_index = first_index - 1 #ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                            
                            #Get the last index in the range
                            last_index_start = run.text.find(":") + 1
                            last_index_end = run.text.find(specifier_end, last_index_start)
                            last_index = run.text[last_index_start:last_index_end]
                            OG_last_index = "null"
                            if last_index == "~":
                                OG_last_index = last_index
                                last_index = len(export_value_array)
                            if last_index == "-~":
                                OG_last_index = last_index
                                last_index = 1
                            last_index = int(last_index)
                            last_index = last_index - 1#ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                            
                            if first_index == last_index: #the range is a single digit long(example: 1:1, 2:2, etc.)
                                #replace the key with the single specified value, and do so while maintaing the user's format placeholders  
                                index = first_index
                                value = getattr(obj, export_key)[index]
                                key_txt = import_key + "[#" + str(index + 1) + ":" + str(index + 1) + "]"
                                run.text = run.text.replace(key_txt, value)
                            else: #the range is muiltiple digits
                                back_to_front = False
                                if first_index > last_index: #check if the user wants the values listed from back to front
                                    back_to_front = True 
                                #replace the key with each value individually and compile an array of these replacements
                                #this will ensure the user's formatting placeholders aren't lost during import
                                import_value_array = []
                                if OG_first_index != "null" and OG_last_index == "null":
                                    key_txt = import_key + "[#" + OG_first_index + ":" + str(last_index + 1) + "]"
                                elif OG_last_index != "null" and OG_last_index == "null":
                                    key_txt = import_key + "[#" + str(first_index + 1) + ":" + OG_last_index + "]"
                                elif OG_first_index != "null" and OG_last_index != "null":
                                    key_txt = import_key + "[#" + OG_first_index + ":" + OG_last_index + "]"
                                else:
                                    key_txt = import_key + "[#" + str(first_index + 1) + ":" + str(last_index + 1) + "]"
                                    
                                if back_to_front == True: #format the list of values back to front
                                    index = last_index
                                    while index <= first_index:
                                        import_value_array.append(run.text.replace(key_txt, export_value_array[index]))
                                        index += 1
                                    value_index = len(import_value_array) - 1
                                    while value_index != -1:
                                        if import_value_array[value_index] == import_value_array[len(import_value_array) - 1]:
                                            run.text = run.text.replace(run.text, import_value_array[value_index])
                                        else:
                                            run.text = run.text + import_value_array[value_index]
                                        value_index -= 1
                                else: #format the list of values front to back
                                    index = first_index
                                    while index <= last_index:
                                        import_value_array.append(run.text.replace(key_txt, export_value_array[index]))
                                        index += 1
                                    for value in import_value_array:
                                        if value == import_value_array[0]:
                                            run.text = run.text.replace(run.text, value)
                                        else:
                                            run.text = run.text + value    
                        elif "~" in index_specifier: #The index specifier is requesting the entire array of values
                            export_value_array = getattr(obj, export_key) #export the value array that was saved to the object 
                            value_index = 0
                            import_value_array = []
                            if "[#-~]" in run.text: #check if the user wants the list structured back to front
                                key_txt = import_key + "[#-~]"
                                for value in export_value_array:
                                    import_value_array.append(run.text.replace(key_txt, value))
                                    value_index += 1
                                value_index = len(import_value_array) - 1
                                while value_index != -1:
                                    if import_value_array[value_index] == import_value_array[len(import_value_array) - 1]:
                                        run.text = run.text.replace(run.text, import_value_array[value_index])
                                    else:
                                        run.text = run.text + import_value_array[value_index]
                                    value_index -= 1
                            else: #default is front to back
                                key_txt = import_key + "[#~]"
                                #replace the key with each value individually and compile an array of these replacements
                                #this will ensure the user's formatting placeholders aren't lost during import
                                for value in export_value_array:
                                    import_value_array.append(run.text.replace(key_txt, export_value_array[value_index])) 
                                    value_index += 1
                                #replace the key with the an iterated list of values, with each value wrapped in the user specified formatting placeholders 
                                for value in import_value_array:
                                    if value == import_value_array[0]:
                                        run.text = run.text.replace(run.text, value)
                                    else:
                                        run.text = run.text + value
                        else: #The index specifier is a single index
                            #initialize the import value array
                            import_value_array = []
                            #Get the index 
                            index_start = run.text.find("#") + 1
                            index_end = run.text.find(specifier_end, index_start)
                            index = run.text[index_start:index_end]
                            index = int(index) - 1 #ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                            #replace the key with the single specified value, and do so while maintaing the user's format placeholders  
                            value = getattr(obj, export_key)[index]
                            key_txt = import_key + "[#" + str(index + 1) + "]"
                            import_value_array.append(run.text.replace(key_txt, value))
                            run.text = run.text.replace(key_txt, value)
                    else: #There is no index specifier, the default is the first index
                        #initialize the import value array
                        import_value_array = []
                        value = getattr(obj, export_key)[0]
                        key_txt = import_key
                        import_value_array.append(run.text.replace(key_txt, value))
                        run.text = run.text.replace(key_txt, value)
                    #check for conditional formatting, apply the formatting only to the values that meet the condition
                    for conditional_specifier_type in conditional_specifiers:
                        #Obtain an array that seperates each value from one another, making it easy to modify them individually
                        value_array = import_value_array #Since the import value array already does this much, just grab it and reuse it for a new purpose 
                        if conditional_specifier_type in run.text:
                            if_amount = value_array[0].count(conditional_specifier_type)#get the total amount of if statements of this type that appears in a singular value
                            if_index = 1
                            while if_index <= if_amount:#iterate through the following code until each instance of the if statement type is evaluated
                                if conditional_specifier_type == "if(first)":
                                    #get the conditional specifier, then extract the conditional instruction from it
                                    conditional_specifier_start = run.text.find(conditional_specifier_type) - 1
                                    conditional_specifier_end = run.text.find(";", conditional_specifier_start) + 1 + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    conditional_instruction_start = conditional_specifier.find(")") + 1
                                    conditional_instruction_end = conditional_specifier.find(";")
                                    conditional_instruction = conditional_specifier[conditional_instruction_start:conditional_instruction_end]
                                    
                                    #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                    #on the second iteration: update the value array to reflect the modified values 
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        while value_index < len(value_array):
                                            if value_index == 0:
                                                value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            else:
                                                value = value_array[value_index].replace(conditional_specifier, "")
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1 
                                elif conditional_specifier_type == "if(-last)":
                                    #get the conditional specifier, then extract the conditional instruction from it
                                    conditional_specifier_start = run.text.find(conditional_specifier_type) - 1
                                    conditional_specifier_end = run.text.find(";", conditional_specifier_start) + 1 + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    conditional_instruction_start = conditional_specifier.find(")") + 1
                                    conditional_instruction_end = conditional_specifier.find(";")
                                    conditional_instruction = conditional_specifier[conditional_instruction_start:conditional_instruction_end]
                                    
                                    #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                    #on the second iteration: update the value array to reflect the modified values 
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        last_value_index = len(value_array) - 1
                                        while value_index < len(value_array):
                                            if value_index != last_value_index:
                                                value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            else:
                                                value = value_array[value_index].replace(conditional_specifier, "")
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                                    
                                elif conditional_specifier_type == "if(last)":
                                    #get the conditional specifier, then extract the conditional instruction from it
                                    conditional_specifier_start = run.text.find(conditional_specifier_type) - 1
                                    conditional_specifier_end = run.text.find(";", conditional_specifier_start) + 1 + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    conditional_instruction_start = conditional_specifier.find(")") + 1
                                    conditional_instruction_end = conditional_specifier.find(";")
                                    conditional_instruction = conditional_specifier[conditional_instruction_start:conditional_instruction_end]
                                    
                                    #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                    #on the second iteration: update the value array to reflect the modified values 
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        last_value_index = len(value_array) - 1
                                        while value_index < len(value_array):
                                            if value_index == last_value_index:
                                                value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            else:
                                                value = value_array[value_index].replace(conditional_specifier, "")
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                                elif conditional_specifier_type == "if(-first)":
                                    #get the conditional specifier, then extract the conditional instruction from it
                                    conditional_specifier_start = run.text.find(conditional_specifier_type) - 1
                                    conditional_specifier_end = run.text.find(";", conditional_specifier_start) + 1 + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    conditional_instruction_start = conditional_specifier.find(")") + 1
                                    conditional_instruction_end = conditional_specifier.find(";")
                                    conditional_instruction = conditional_specifier[conditional_instruction_start:conditional_instruction_end]
                                    
                                    #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                    #on the second iteration: update the value array to reflect the modified values 
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        while value_index < len(value_array):
                                            if value_index != 0:
                                                value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            else:
                                                value = value_array[value_index].replace(conditional_specifier, "")
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                                if_index += 1
                    #check for advanced conditional formatting, apply the formatting only to the values that meet the condition
                    if "if(#" in run.text:
                        if_amount = value_array[0].count("if(#")#Get the total amount of advanced if statements that appear in a singular value
                        if_iteration = 1
                        while if_iteration <= if_amount: #Iterate through the following code until each if statement has been evaluated 
                            if_start = run.text.find("if(#")
                            if_end = run.text.find(")", if_start) + 1
                            condition = run.text[if_start:if_end]
                            if ":" in condition:
                                #Get the first index in the range
                                first_index_start = if_start + len("if(#")
                                first_index_end = run.text.find(":", first_index_start)
                                first_index = run.text[first_index_start:first_index_end]
                                infinity_used_first = False
                                try:
                                    first_index = int(first_index)   
                                    first_index = first_index - 1 #ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                                except ValueError:
                                    if first_index == "-~":
                                        first_index = 0
                                        infinity_used_first = True
                                    else:
                                        print("ERROR: the first index of an if statement is not a number")
                                        return "null"
                                #Get the last index in the range
                                last_index_start = run.text.find(":") + 1
                                last_index_end = run.text.find(")", last_index_start)
                                last_index = run.text[last_index_start:last_index_end]
                                infinity_used_last = False
                                try:
                                    last_index = int(last_index)
                                    last_index = last_index - 1#ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                                except ValueError:
                                    if last_index == "~":
                                        last_index = len(value_array) - 1
                                        infinity_used_last = True
                                    else:
                                        print("ERROR: the last index of an if statement is not a number")
                                        return "null"  
                                #Get the conditional instruction
                                instruction_start = run.text.find(")[") + 1
                                instruction_end = run.text.find("];") + 1
                                conditional_instruction = run.text[instruction_start:instruction_end]
                                
                                #Get the conditional specifier
                                if infinity_used_first == False and infinity_used_last == False:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(#" + str(first_index + 1) + ":" + str(last_index + 1) + ")" + conditional_instruction + ";" + specifier_end)
                                elif infinity_used_first == True and infinity_used_last == False:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(#" + "-~" + ":" + str(last_index + 1) + ")" + conditional_instruction + ";" + specifier_end)
                                elif infinity_used_last == True and infinity_used_first == False:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(#" + str(first_index + 1) + ":" + "~" + ")" + conditional_instruction + ";" + specifier_end)
                                elif infinity_used_first == True and infinity_used_last == True:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(#" + "-~" + ":" + "~" + ")" + conditional_instruction + ";" + specifier_end)
                                conditional_specifier_end = run.text.find(";" + specifier_end, conditional_specifier_start) + len(";") + len(specifier_end)
                                conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                
                                #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                #on the second iteration: update the value array to reflect the modified values 
                                iteration = 0
                                while iteration != 2:
                                    value_index = 0
                                    while value_index < len(value_array):
                                        if value_index >= first_index and value_index <= last_index:
                                            value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                            if iteration == 1:
                                                value_array[value_index] = value
                                        else:
                                            value = value_array[value_index].replace(conditional_specifier, "")
                                            if iteration == 1:
                                                value_array[value_index] = value
                                        if iteration == 0:
                                            run.text = run.text.replace(value_array[value_index], value)
                                        value_index += 1
                                    iteration += 1      
                            else:
                                #Get the first index in the range
                                index_start = if_start + len("if(#")
                                index_end = run.text.find(")", index_start)
                                index = run.text[index_start:index_end]
                                if index == "~":
                                    #Get the conditional instruction
                                    instruction_start = run.text.find(")[") + 1
                                    instruction_end = run.text.find("];") + 1
                                    conditional_instruction = run.text[instruction_start:instruction_end]
                                    
                                    #Get the conditional specifier
                                    conditional_specifier_start = run.text.find(specifier_start + "if(#" + "~" + ")" + conditional_instruction + ";" + specifier_end)
                                    conditional_specifier_end = run.text.find(";" + specifier_end, conditional_specifier_start) + len(";") + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        while value_index < len(value_array):
                                            value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                            if iteration == 1:
                                                value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                                else:
                                    index = int(index)
                                    index = index - 1 #ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                                   
                                    #Get the conditional instruction
                                    instruction_start = run.text.find(")[") + 1
                                    instruction_end = run.text.find("];") + 1
                                    conditional_instruction = run.text[instruction_start:instruction_end]
                                    
                                    #Get the conditional specifier
                                    conditional_specifier_start = run.text.find(specifier_start + "if(#" + str(index + 1) + ")" + conditional_instruction + ";" + specifier_end)
                                    conditional_specifier_end = run.text.find(";" + specifier_end, conditional_specifier_start) + len(";") + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    
                                    #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                    #on the second iteration: update the value array to reflect the modified values 
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        while value_index < len(value_array):
                                            if value_index == index:
                                                value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            else:
                                                value = value_array[value_index].replace(conditional_specifier, "")
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                            if_iteration += 1
                    
                    #IF NOT conditional
                    if "if(-#" in run.text:
                        if_amount = value_array[0].count("if(-#")#Get the total amount of advanced if statements that appear in a singular value
                        if_iteration = 1
                        while if_iteration <= if_amount: #Iterate through the following code until each if statement has been evaluated 
                            if_start = run.text.find("if(-#")
                            if_end = run.text.find(")", if_start) + 1
                            condition = run.text[if_start:if_end]
                            if ":" in condition:
                                #Get the first index in the range
                                first_index_start = if_start + len("if(-#")
                                first_index_end = run.text.find(":", first_index_start)
                                first_index = run.text[first_index_start:first_index_end]
                                infinity_used_first = False
                                try:
                                    first_index = int(first_index)   
                                    first_index = first_index - 1 #ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                                except ValueError:
                                    if first_index == "-~":
                                        first_index = 0
                                        infinity_used_first = True
                                    else:
                                        print("ERROR: the first index of an if statement is not a number")
                                        return "null"
                                #Get the last index in the range
                                last_index_start = run.text.find(":") + 1
                                last_index_end = run.text.find(")", last_index_start)
                                last_index = run.text[last_index_start:last_index_end]
                                infinity_used_last = False
                                try:
                                    last_index = int(last_index)
                                    last_index = last_index - 1#ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                                except ValueError:
                                    if last_index == "~":
                                        last_index = len(value_array) - 1
                                        infinity_used_last = True
                                    else:
                                        print("ERROR: the last index of an if statement is not a number")
                                        return "null"  
                                
                                #Get the conditional instruction
                                instruction_start = run.text.find(")[") + 1
                                instruction_end = run.text.find("];") + 1
                                conditional_instruction = run.text[instruction_start:instruction_end]
                                
                                #Get the conditional specifier
                                if infinity_used_first == False and infinity_used_last == False:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(-#" + str(first_index + 1) + ":" + str(last_index + 1) + ")" + conditional_instruction + ";" + specifier_end)
                                elif infinity_used_first == True and infinity_used_last == False:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(-#" + "-~" + ":" + str(last_index + 1) + ")" + conditional_instruction + ";" + specifier_end)
                                elif infinity_used_last == True and infinity_used_first == False:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(-#" + str(first_index + 1) + ":" + "~" + ")" + conditional_instruction + ";" + specifier_end)
                                elif infinity_used_first == True and infinity_used_last == True:
                                    conditional_specifier_start = run.text.find(specifier_start + "if(-#" + "-~" + ":" + "~" + ")" + conditional_instruction + ";" + specifier_end)
                                conditional_specifier_end = run.text.find(";" + specifier_end, conditional_specifier_start) + len(";") + len(specifier_end)
                                conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                
                                #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                #on the second iteration: update the value array to reflect the modified values 
                                iteration = 0
                                while iteration != 2:
                                    value_index = 0
                                    while value_index < len(value_array):
                                        if value_index < first_index or value_index > last_index:
                                            value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                            if iteration == 1:
                                                value_array[value_index] = value
                                        else:
                                            value = value_array[value_index].replace(conditional_specifier, "")
                                            if iteration == 1:
                                                value_array[value_index] = value
                                        if iteration == 0:
                                            run.text = run.text.replace(value_array[value_index], value)
                                        value_index += 1
                                    iteration += 1      
                            else:
                                #Get the first index in the range
                                index_start = if_start + len("if(-#")
                                index_end = run.text.find(")", index_start)
                                index = run.text[index_start:index_end]
                                if index == "~":
                                    #Get the conditional instruction
                                    instruction_start = run.text.find(")[") + 1
                                    instruction_end = run.text.find("];") + 1
                                    conditional_instruction = run.text[instruction_start:instruction_end]
                                    
                                    #Get the conditional specifier
                                    conditional_specifier_start = run.text.find(specifier_start + "if(-#" + "~" + ")" + conditional_instruction + ";" + specifier_end)
                                    conditional_specifier_end = run.text.find(";" + specifier_end, conditional_specifier_start) + len(";") + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        while value_index < len(value_array):
                                            value = value_array[value_index].replace(conditional_specifier, "")
                                            if iteration == 1:
                                                value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                                else:
                                    index = int(index)
                                    index = index - 1 #ensures that the value array starts at 0, as opposed to 1(since users will be not be expected to conform to the formatting of coding arrays)
                                   
                                    #Get the conditional instruction
                                    instruction_start = run.text.find(")[") + 1
                                    instruction_end = run.text.find("];") + 1
                                    conditional_instruction = run.text[instruction_start:instruction_end]
                                    
                                    #Get the conditional specifier
                                    conditional_specifier_start = run.text.find(specifier_start + "if(-#" + str(index + 1) + ")" + conditional_instruction + ";" + specifier_end)
                                    conditional_specifier_end = run.text.find(";" + specifier_end, conditional_specifier_start) + len(";") + len(specifier_end)
                                    conditional_specifier = run.text[conditional_specifier_start:conditional_specifier_end]
                                    
                                    #on the first iteration: if the value meets the condition, then keep the conditional formatting specifiers, if it does not, remove them 
                                    #on the second iteration: update the value array to reflect the modified values 
                                    iteration = 0
                                    while iteration != 2:
                                        value_index = 0
                                        while value_index < len(value_array):
                                            if value_index != index:
                                                value = value_array[value_index].replace(conditional_specifier, conditional_instruction)
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            else:
                                                value = value_array[value_index].replace(conditional_specifier, "")
                                                if iteration == 1:
                                                    value_array[value_index] = value
                                            if iteration == 0:
                                                run.text = run.text.replace(value_array[value_index], value)
                                            value_index += 1
                                        iteration += 1
                            if_iteration += 1
                 
                    #check for custom list placeholders
                    left_smart_quote_ASCII = 8220
                    right_smart_quote_ASCII = 8221
                    left_smart_quote = chr(left_smart_quote_ASCII)
                    right_smart_quote = chr(right_smart_quote_ASCII)
                    custom_list_left_indicator = "[[" + left_smart_quote
                    custom_list_right_indicator = right_smart_quote + "]]"
                    if custom_list_left_indicator in run.text and custom_list_right_indicator in run.text:
                        custom_list_start = value_array[0].find(custom_list_left_indicator)
                        custom_list_end = value_array[0].find(custom_list_right_indicator, custom_list_start) + len(custom_list_right_indicator)
                        custom_list = value_array[0][custom_list_start:custom_list_end]
                        list_member_amount = custom_list.count(left_smart_quote)
                        list_member_index = 1
                        list_member_start = 0
                        list_member_end = 0
                        list_members = []
                        while list_member_index <= list_member_amount:
                            list_member_start = custom_list.find(left_smart_quote, list_member_end)
                            list_member_end = custom_list.find(right_smart_quote, list_member_start)
                            list_member = custom_list[list_member_start + len(left_smart_quote):list_member_end]
                            list_member_index += 1
                            list_members.append(list_member)
                        list_member_index = 0
                        value_index = 0
                        for value in value_array:
                            if list_member_index == len(list_members):
                                list_member_index = 0
                            value = value.replace(custom_list, list_members[list_member_index])
                            run.text = run.text.replace(value_array[value_index], value)
                            value_array[value_index] = value
                            value_index += 1
                            list_member_index += 1
                    
                    #check for custom list placeholders
                    left_smart_quote_ASCII = 8220
                    right_smart_quote_ASCII = 8221
                    left_smart_quote = chr(left_smart_quote_ASCII)
                    right_smart_quote = chr(right_smart_quote_ASCII)
                    custom_list_left_indicator = "[-[" + left_smart_quote
                    custom_list_right_indicator = right_smart_quote + "]]"
                    if custom_list_left_indicator in run.text and custom_list_right_indicator in run.text:
                        custom_list_start = value_array[0].find(custom_list_left_indicator)
                        custom_list_end = value_array[0].find(custom_list_right_indicator, custom_list_start) + len(custom_list_right_indicator)
                        custom_list = value_array[0][custom_list_start:custom_list_end]
                        list_member_amount = custom_list.count(left_smart_quote)
                        list_member_index = 1
                        list_member_start = 0
                        list_member_end = 0
                        list_members = []
                        while list_member_index <= list_member_amount:
                            list_member_start = custom_list.find(left_smart_quote, list_member_end)
                            list_member_end = custom_list.find(right_smart_quote, list_member_start)
                            list_member = custom_list[list_member_start + len(left_smart_quote):list_member_end]
                            list_member_index += 1
                            list_members.append(list_member)
                        list_member_index = len(list_members) - 1
                        value_index = 0
                        for value in value_array:
                            if list_member_index == -1:
                                list_member_index = len(list_members) - 1
                            value = value.replace(custom_list, list_members[list_member_index])
                            run.text = run.text.replace(value_array[value_index], value)
                            value_array[value_index] = value
                            value_index += 1
                            list_member_index -= 1
                            
                    #check for string placeholders 
                    if left_smart_quote in run.text and right_smart_quote in run.text:
                        run.text = run.text.replace("[" + left_smart_quote, "")
                        run.text = run.text.replace(right_smart_quote + "]", "")
     
                    #replace all instances of non-character place holders with their corresponding formatting 
                    for place_holder in place_holders: 
                        if place_holder[0] in run.text:
                            run.text = run.text.replace(specifier_start + place_holder[0] + specifier_end, place_holder[1])
                    
                    #if there is a list placeholder, pair each one of the object's values with an equivelant value retrieved from the placeholder's corresponding array 
                    alphabet_array = ["a", "b", "c", "d", "e", "f", "g", "h", "i", "j", "k", "l", "m", "n", "o", "p", "r", "s", "t", "u", "v", "w", "x", "y", "z"]
                    ALPHABET_array = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "R", "S", "T", "U", "V", "W", "X", "Y", "Z"]
                    xyz_array = ["x", "y", "z"]
                    try:
                        number_array = list(range(1, len(import_value_array) + 1))
                    except UnboundLocalError:
                        number_array = [1]
                    for list_place_holder in list_place_holders:
                        if list_place_holder in run.text:
                            #pattern: x -> y -> z -> x -> y -> z, etc.
                            if list_place_holder == "xyz_index":
                                xyz_index = 0
                                value_index = 0
                                for value in value_array:
                                    if xyz_index == len(xyz_array):
                                        xyz_index = 0
                                    value = value.replace("[" + list_place_holder + "]", xyz_array[xyz_index])
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    xyz_index += 1
                                    value_index += 1
                            #pattern: z -> y -> x -> z -> y -> x, etc.
                            elif list_place_holder == "-xyz_index":
                                xyz_index = len(xyz_array) - 1
                                value_index = 0
                                for value in value_array:
                                    if xyz_index == -1:
                                        xyz_index = len(xyz_array) - 1
                                    value = value.replace("[" + list_place_holder + "]", xyz_array[xyz_index])
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    xyz_index -= 1
                                    value_index += 1
                            #pattern: a -> b -> c -> ... -> z -> a -> b -> c -> ... -> z, etc.
                            elif list_place_holder == "alphabet_index":
                                alphabet_index = 0
                                value_index = 0
                                for value in value_array:
                                    if alphabet_index == len(alphabet_array):
                                        alphabet_index = 0
                                    value = value.replace("[" + list_place_holder + "]", alphabet_array[alphabet_index])
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    alphabet_index += 1
                                    value_index += 1
                            #pattern: z -> y -> x -> ... -> a -> z -> y -> x -> ... -> a, etc.
                            elif list_place_holder == "-alphabet_index":
                                alphabet_index = len(alphabet_array) - 1
                                value_index = 0
                                for value in value_array:
                                    if alphabet_index == -1:
                                        alphabet_index = len(alphabet_array) - 1
                                    value = value.replace("[" + list_place_holder + "]", alphabet_array[alphabet_index])
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    alphabet_index -= 1
                                    value_index += 1
                            #pattern: A -> B -> C -> ... -> Z -> A -> B -> C -> ... -> A, etc.
                            elif list_place_holder == "ALPHABET_index":
                                alphabet_index = 0
                                value_index = 0
                                for value in value_array:
                                    if alphabet_index == len(ALPHABET_array):
                                        alphabet_index = 0
                                    value = value.replace("[" + list_place_holder + "]", ALPHABET_array[alphabet_index])
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    alphabet_index += 1
                                    value_index += 1
                            #pattern: Z -> Y -> X -> ... -> A -> Z -> Y -> X -> ... -> A, etc.
                            elif list_place_holder == "-ALPHABET_index":
                                alphabet_index = len(ALPHABET_array) - 1
                                value_index = 0
                                for value in value_array:
                                    if alphabet_index == -1:
                                        alphabet_index = len(ALPHABET_array) - 1
                                    value = value.replace("[" + list_place_holder + "]", ALPHABET_array[alphabet_index])
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    alphabet_index -= 1
                                    value_index += 1
                            #pattern: 1 -> 2 -> 3, etc.
                            elif list_place_holder == "number_index":
                                number_index = 0
                                value_index = 0
                                for value in value_array:
                                    if number_index == len(number_array):
                                        number_index = 0
                                    value = value.replace("[" + list_place_holder + "]", str(number_array[number_index]))
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    number_index += 1
                                    value_index += 1
                            #pattern: 3 -> 2 -> 1, etc.
                            elif list_place_holder == "-number_index":
                                number_index = len(number_array) - 1
                                value_index = 0
                                for value in value_array:
                                    if number_index == -1:
                                        number_index = len(number_array) - 1
                                    value = value.replace("[" + list_place_holder + "]", str(number_array[number_index]))
                                    run.text = run.text.replace(value_array[value_index], value)
                                    value_array[value_index] = value 
                                    number_index -= 1
                                    value_index += 1
                    #remove the leftover key indicators 
                    run.text = run.text.replace(key_start_indicator, "")
                    run.text = run.text.replace(key_end_indicator, "")

    #Save the export 
    doc.save(doc_path)
    print("value or values exported to all instances of the following key: " + obj_type + "." + key)

def ToDigit(character_string):
    '''
    Transforms a string of text into a string of digits. Each pair of digits is an arbritary value associated with a given letter or non-letter character. Non-letter characters are manually replaced with their chosen equivelant, while letter characters are automatically replaced by looping through the ascii_letters array. 
    '''
    
    #Replaces all instances of these non-letter characters with their digit equivelant
    character_string = character_string.replace(" ", "00")
    character_string = character_string.replace("-", "53")
    character_string = character_string.replace("--", "54")
    
    #Replaces all instances of a given letter with their digit equivelant 
    digit = 0
    for letter in ascii_letters:
        digit += 1
        if digit < 10:
            character_string = character_string.replace(letter, "0" + str(digit))
        else:
            character_string = character_string.replace(letter, str(digit))    
    digit_string = character_string
    
    #Returns the string of digits
    return digit_string
      
def ToCharacter(digit_string):
    '''
    Transforms a string of digits into a string of text. Each digit pair is transformed into its corresponding letter or non-letter character. Non-letter character digit pairs are manually transformed into their equivelants. Letter character digit pairs are automatically transformed into their equivelants by looping through the ascii letters array. 
    '''
    digit_string = str(digit_string)
    digit_length = len(digit_string) * .5 #The length is halved to represent the amount of two-digit pairs
    digit_length = int(digit_length)
    
    #Creates a two-dimensional array that pairs a given letter with their digit equivelant by looping through the ascii letter array
    letters = ["null"] #Index 0 is declared null so each digit pair, which starts at 01, lines up with its corresponding letter array index 
    digit_equivelant = 0
    for letter in ascii_letters:
        digit_equivelant += 1
        if digit_equivelant < 10:
            letters.append(["0" + str(digit_equivelant), letter])
        else:
            letters.append([digit_equivelant, letter])
    
    #Creates a string of text, wherein each digit has been replaced with its letter or non-letter equivelant 
    txt = ""
    digit_pair_index = 0
    digit_pair_start = 0
    digit_pair_end = 2
    while digit_pair_index < digit_length:
        digit_pair = digit_string[digit_pair_start:digit_pair_end]

        if digit_pair == "00":
            letter = " "
        elif digit_pair == "53":
            letter = "-"
        elif digit_pair == "54":
            letter = "--"
        else:
            digit_pair = int(digit_pair)
            letter = letters[digit_pair][1] 
        txt += letter
       
        digit_pair_index += 1
        digit_pair_start += 2
        digit_pair_end += 2  
    return txt
    
def GenerateID(obj):
    '''
    This functions generates an ID. Each ID consists of two values: the object type and the object name. The object type is a two digit string that precedes the object name. If the object provided is not a valid object type, then the ID is returned as null. 
    '''
    #Get the type tag associated with the object type, if the object type does not exist, then the type tag and the ID is null
    ID = "null"
    type_tag = "null"
    obj_type = GetObjType(obj, "null")
    if obj_type == "DocType":
        type_tag = 0    #elif_end
    elif obj_type == "Doc":
        type_tag = 1    #elif_end
    elif obj_type == "Lawyer":
        type_tag = 2    #elif_end
    elif obj_type == "OC":
        type_tag = 3    #elif_end
    elif obj_type == "Employer":
        type_tag = 4    #elif_end
    elif obj_type == "Carrier":
        type_tag = 5    #elif_end
    elif obj_type == "ServicingAgent":
        type_tag = 6    #elif_end 
    elif obj_type == "Adjuster":
        type_tag = 7    #elif_end 
    elif obj_type == "EmployerRep":
        type_tag = 8    #elif_end 
    elif obj_type == "Entity":
        type_tag = 9    #elif_end 
    elif obj_type == "Claimant":
        type_tag = 10   #elif_end 
    elif obj_type == "SpecialLanguage":
        type_tag = 11   #elif_end 
    elif obj_type == "RegisteredAgent":
        type_tag = 12   #elif_end 
  #Ensures that null is returned when the object type is invalid 
    if type_tag == "null":
        ID = "null"
        return ID
    else:
        ID = str(type_tag) + ToDigit(getattr(obj, "obj_name")[0])
    #Ensures that the type tag is always two digits
    if type_tag < 10:
        ID = "0" + ID
    return ID
    #END    
def ArrayID(ID):
    '''Retrieves each digit pair that comprises a given ID, and then organizes them into a neatly seperated array'''
    
    #Iterates through each digit pair in the ID, and appends the digit pair to the digit pairs array
    ID = str(ID)
    ID_length = len(ID) * .5 #The length is halved to represent the amount of two-digit pairs
    digit_pair_start = 0
    digit_pair_end = 2
    digit_pairs = []
    digit_pair_index = 0
    while digit_pair_index < len(ID):
        digit_pair = ID[digit_pair_start:digit_pair_end]
        digit_pairs.append(digit_pair)
        digit_pair_start += 2
        digit_pair_end += 2
        digit_pair_index += 2
    return digit_pairs

def CreateEmptyObj(obj_type):
    '''Returns an empty object of the desired type'''
    obj = "null"
    if obj_type == "Doc":
        obj = Doc()
    elif obj_type == "MuiltiDoc":
        obj = MuiltiDoc()
    elif obj_type == "Lawyer":
        obj = Lawyer()
    elif obj_type == "OC":
        obj = OC()
    elif obj_type == "Employer":
        obj = Employer()
    elif obj_type == "Carrier":
        obj = Carrier()
    elif obj_type == "ServicingAgent":
        obj = ServicingAgent()
    elif obj_type == "Adjuster":
        obj = Adjuster()
    elif obj_type == "EmployerRep":
        obj = EmployerRep()
    elif obj_type == "Entity":
        obj = Entity()
    elif obj_type == "Claimant":
        obj = Claimant()
    elif obj_type == "SpecialLanguage":
        obj = SpecialLanguage()
    elif obj_type == "RegisteredAgent":
        obj = RegisteredAgent()
    return obj
    #END  
def GetObjKeys(ID = "null", obj = "null", obj_type = "null"):
    #Get the object from the object type 
    if obj != "null":
        obj_type = GetObjType(obj)
        obj = CreateObj(obj_type)
    elif ID != "null":
        obj_type = GetObjType(ID)
        obj = CreateObj(obj_type)
    elif obj_type != "null":
        obj = CreateEmptyObj(obj_type)
    

    #Get the keys associated with an object, and store them as a single cleaned up string
    keys = str(obj.__dict__.keys())
    keys = keys.replace("dict_keys", "")
    keys = keys.replace("(", "")
    keys = keys.replace(")", "")
    keys = keys.replace("[", "")
    keys = keys.replace("]", "")
    keys = keys.replace("'", "")
    
    #Deconstruct the string of keys into an array of keys 
    key_start = 0
    key_end = 0
    skip_identifier = len(", ")
    
    key_index = 0
    key = "null"
    key_array = []
    while key_index < keys.count(", ") + 1:
        if key_index != 0:
            key_start = keys.find(", ", key_end)
                        
        if key_start != 0:
            key_end = keys.find(", ", key_start + skip_identifier)
        else:
            key_end = keys.find(", ", key_start)
                        
        if key_end == -1:
            key_end = len(keys)
                        
        if key_index != 0:
            key = keys[key_start + skip_identifier:key_end]
        else:
            key = keys[key_start:key_end]
        key_array.append(key)
        key_index += 1
    return key_array  

def GetKeyColumns(obj_type, obj_type_keys = "null"):
    '''
    Using the amount of keys associated with an object type, return an array with an equal amount of columns. 
    Each key is correlated to a column by index, for example, the first key of an object type is correlated to the first column of the spreadsheet.
    (starting at column B, since the A column is always designated to IDs)
    '''
    
    #Get the keys of a given object type and then get  thecolumns to store them, or get the columns that are already storing them, as both are the same
    if obj_type_keys == "null":
        obj_type_keys = GetObjKeys("null", "null", obj_type)
    key_amount = len(obj_type_keys)
    letters = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q", "R", "S", "T", "U", "V", "W", "X", "Y", "Z"]
    columns = []
    column_array_complete = False
    while column_array_complete == False: 
        for letter in letters:
            if letter != "A":
                columns.append(letter)
            column_amount = len(columns)
            if column_amount == key_amount:
                break       
        if column_amount == key_amount:
            break
            
        for first_letter in letters:
            if column_amount == key_amount:
                break
            for second_letter in letters:
                columns.append(first_letter + second_letter)
                column_amount = len(columns)
                if column_amount == key_amount:
                    break
        if column_amount == key_amount:
            break
        
        last_column = "null" #At this stage, the last column must be checked because the last column in an excel spreadsheet is XFD
        for first_letter in letters:
            if column_amount == key_amount or last_column == "XFD":
                break
            for second_letter in letters:
                if column_amount == key_amount or last_column == "XFD":
                    break
                for third_letter in letters:
                    columns.append(first_letter + second_letter + third_letter)
                    column_amount = len(columns)
                    last_column_index = column_amount - 1
                    last_column = columns[last_column_index]
                    if column_amount == key_amount or last_column == "XFD":
                        break
        if column_amount != key_amount:
            print("Error: There are more than 16,384 keys for this object type, which exceeds the amount of rows available in an excel spreadsheet.\nAny key past the 16,384th key has not been saved.\nPlease shorten the amount of keys associated with this object type for this object type to work as intended.")
        if column_amount == key_amount or last_column == "XFD1":
            break
    return columns
    
def GetObjType(obj = "null", ID = "null"):
    '''
    Get an object type by using either an ID or an object.
    If the ID is provided, identify the object type using the ID's two digit type tag, otherwise simply get the object's name
    '''
    
    if obj != "null":
        obj_type = str(type(obj).__name__)
    elif ID != "null":
        obj_type = "null"
        type_tag = str(ID[0:2])
        obj_type_identified = False
        if type_tag == "00":
            obj_type = "Doc"#elif_end       
        elif type_tag == "01":
            obj_type = "MuiltiDoc"#elif_end 
        elif type_tag == "02":
            obj_type = "Lawyer"#elif_end
        elif type_tag == "03":
            obj_type = "OC"#elif_end 
        elif type_tag == "04":
            obj_type = "Employer"#elif_end
        elif type_tag == "05":
            obj_type = "Carrier"#elif_end
        elif type_tag == "06":
            obj_type = "ServicingAgent"#elif_end
        elif type_tag == "07":
            obj_type = "Adjuster"#elif_end
        elif type_tag == "08":
            obj_type = "EmployerRep"#elif_end
        elif type_tag == "09":
            obj_type = "Entity"#elif_end
        elif type_tag == "10":
            obj_type = "Claimant"#elif_end
        elif type_tag == "11":
            obj_type = "SpecialLanguage"#elif_end
        elif type_tag == "12":
            obj_type = "RegisteredAgent"#elif_end 
       
    return obj_type
    #END    
def CreateSaveFile(obj_type, obj_type_keys = "null"):
    '''
    This function creates a new save file or overwrites an existing save file for a given object type by: 
    creating a new excel spreadsheet via openpyxl, creating an ID column, and lastly assigning each of the object type's keys a column
    '''
    #Create the excel spreadsheet to store data pertaining to an object's type
    save_file_name = obj_type.lower() + "_data"
    save_file_PATH = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "data\\" + save_file_name + ".xlsx")
    save_file = Workbook()
    save_file.save(save_file_PATH)
    
    #Add the ID column to the spreadsheet 
    input_data = save_file.active
    input_data["A1"] = "ID"
    
    #Get an array of columns that can be used to store the keys
    columns = GetKeyColumns(obj_type, obj_type_keys)
    
    #Assign each key a column
    if obj_type_keys == "null":
        obj_type_keys = GetObjKeys("null", "null", obj_type)
    key_index = 0
    for key in obj_type_keys:
        key_column = columns[key_index] + "1" #All the key columns are in the first row of a given spreadsheet 
        input_data[key_column] = key
        key_index += 1
    save_file.save(save_file_PATH)
    print("created a save file for the " + obj_type + " object type")

def GetSaveFile(obj_type):
    '''This function gets an object type's save file by inputting an object type into the standard path for all save files'''
    save_file_name = obj_type.lower() + "_data"
    save_file = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "data\\" + save_file_name + ".xlsx")
    return save_file 

def GetObj(ID):
    '''
    This function uses an ID to retrieve an object from its associated save file. It does this by:
        First, getting the object's type, getting the object's save file using its type, and then preparing the save file to be read
        Second, iterating through the ID column's rows until the ID is found. This row index is the row index for the entire object
        Third, it iterates through the key columns of the row and retrieves the values associated with each
        Fourth, it creates a new object to assign the values to 
        Fifth, it iterates through each value and organizes its sub-values into an array
        Sixth, it creates a new object and then assigns each value array to its associated key
        Seventh, it returns the workable object 
    '''
    obj = "null"
    
    #Get the object type, get the save file for the object type, and then prepare the save file to be read 
    obj_type = GetObjType("null", ID)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    output_data = save_file.active
    
    #Find the ID's row in the ID column and then use it to retrieve the object's row   
    row_index = 2
    ID_cell = "null"
    ID_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    while ID_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if output_data[cell].value == ID:
            ID_cell_found = True
            ID_cell = cell
            obj_row_index = row_index 
            break 
        row_index += 1
    if ID_cell == "null":
        print("Error: object does not exist or object's ID is corrupted")
        return "null"
    
    #First, iterate through the object's keys and retrieve the values associated with each
    #Second, iterate through the value and organize its sub-values into an array by utilizing the indicators that seperate them 
    #Third, assign each value array to its associated key, and then return a workable object once complete 
    key_index = 0
    obj = CreateEmptyObj(obj_type)
    obj_keys = GetObjKeys("null", "null", obj_type)
    columns = GetKeyColumns(obj_type)
    for key in obj_keys:
        key_column = columns[key_index] + str(obj_row_index)
        value = output_data[key_column].value
        if str(value) != "None":
            total_values = value.count(value_start_indicator) - 1 #Index from zero not from one, do this by subtracting one from the total amount of values
            value_index = 0
            last_value_start = 0
            last_value_end = 0
            skip_start_indicator = len(value_start_indicator)
            skip_end_indicator = len(value_end_indicator)
            value_array = []
            while value_index <= total_values:
                value_start = value.find(value_start_indicator, last_value_start)
                value_end = value.find(value_end_indicator, last_value_end)
                value_data = value[value_start + skip_start_indicator:value_end]
                last_value_start = value_start + skip_start_indicator
                last_value_end = value_end + skip_end_indicator
                value_index += 1
                value_array.append(value_data)
            setattr(obj, key, value_array)
        else:
            setattr(obj, key, "")
        key_index += 1
    print("retrieved object: " + obj.obj_name[0])
    return obj

def SaveObj(obj, new_object = False):
    '''
    This function saves an object to its associated save file by:
        First, getting the object's ID and deleting any pre-existing duplicate 
        Second, getting the object's type, keys, and save file path. 
        Third, iterating through the rows of the ID column until it finds an empty row to store the object. 
        Fourth, storing the object's ID in the ID column 
        Fifth, getting the columns associated with each of the object's keys. 
        Sixth, getting the value array associated with each of the object's keys and then organizing the values into an array format fit for a cell
        Seventh, assigning each key's associated value array to the key's associated column.
        Eighth, saving the spreadsheet.
    '''
    #Generate the object's ID
    ID = GenerateID(obj)
    
    #If the object already exists, delete it. Assume that it might unless otherwise stated 
    if new_object == False:
        DeleteObj(ID)
        
    #Get the object type, get the object type's keys, get the save file for the object type, and prepare the save file to be edited 
    obj_type = GetObjType(obj, "null")
    obj_keys = GetObjKeys("null", "null", obj_type)
    ID = GenerateID(obj)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    input_data = save_file.active
    
    #Find the first empty ID cell to store the object's ID. Its row index will indicate the row where the rest of the object should be stored 
    row_index = 2
    empty_ID_cell = "null"
    empty_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    while empty_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if input_data[cell].value == None:
            input_data[cell].number_format = numbers.FORMAT_TEXT #format the ID cell as text to prevent corruption from undefined formatting 
            empty_cell_found = True
            empty_ID_cell = cell
            empty_row_index = row_index 
            break 
        row_index += 1
    if empty_ID_cell == "null":
        print("Error: There is no avaliable cell that can be used to store this object.")
        return "null"
    
    #Store the ID in the first available empty cell in the A column 
    input_data[empty_ID_cell] = ID
    
    #Store the value array in the column associated with each key 
    #IMPORTANT: This algorithm assumes that the key initiations of a given object type's associated class HAVE NOT CHANGED POSITION since the conception of the save file 
    #IMPORTANT: DO NOT add new key iniations BETWEEN existing iniations, as then the code below will NOT work as intended. ONLY add new iniations to the end of a given class
    columns = GetKeyColumns(obj_type)
    column_index = 0
    for key in obj_keys:
        value_array = getattr(obj, key)
        value_array = str(value_array)
        value_array = value_array.replace("[", "")
        value_array = value_array.replace("]", "")
        value_array = value_array.replace("'", "")
        value_array = value_start_indicator + value_array
        value_array = value_array + value_end_indicator
        value_array = value_array.replace(", ", value_end_indicator + value_start_indicator)
        if value_array != value_start_indicator + value_end_indicator: #If the value array is empty, there is no need to waste storage saving it as empty
            cell = columns[column_index] + str(row_index)
            input_data[cell] = value_array
        column_index += 1
    save_file.save(save_file_PATH)
    print("object saved at row " + str(row_index))

def DeleteObj(ID):
    '''
    This function deletes an object by: 
        First, preparing the save file to be modified 
        Second, using an ID to find the object's row
        Third, delete the object's row using excel's built-in deletion ability 
    '''
    
    #Get the object type, get the object type's keys, get the save file for the object type, and prepare the save file to be edited 
    obj_type = GetObjType("null", ID)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    delete_data = save_file.active 
    
    #Find the ID's row in the ID column and then use it to retrieve the object's row   
    row_index = 2
    ID_cell = "null"
    ID_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    while ID_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if delete_data[cell].value == ID:
            ID_cell_found = True
            ID_cell = cell
            obj_row_index = row_index 
            break 
        row_index += 1
    if ID_cell == "null":
        print("Error: object does not exist or object's ID is corrupted")
        return "null"
    delete_data.delete_rows(row_index, 1)
    save_file.save(save_file_PATH)
    print("deleted object at row " + str(row_index))
    
def GetSimilarIDs(ID, similarity_percentage):
    '''
    Get similar IDs within the specified similarity range. Do this by:
        First, preparing the save file to be read
        Second, fetching an ID from a given ID row
        Third, comparing the digit pairs of this comparand ID to the provided ID
        Fourth, dividing the amount of identified similar digit pairs by the amount of total digit pairs
        Fifth, checking to see if this percentage matches the desired similarity percentage 
        Sixth, adding it to the similar IDs array if it does 
        Seventh, iterating to the next row until there are no IDs left to compare 
        Eigth, returning the similar IDs array 
    '''
    #Get the object type, get the save file for the object type, and then prepare the save file to be read 
    obj_type = GetObjType("null", ID)
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    output_data = save_file.active
    
    #Iterate through each row in the ID column until there are no IDs left
    row_index = 2
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    ID_digit_pairs = ArrayID(ID)
    similar_IDs = []
    while row_index != row_max:
        cell = "A" + str(row_index)
        comparand_ID = output_data[cell].value
        if str(comparand_ID) == "None":
            break
        comparand_digit_pairs = ArrayID(comparand_ID) 
        similar_pairs = 0
        #Once a comparand ID has been retrieved compare its digit pairs to the provided IDs digit pairs
        for comparand_digit_pair in comparand_digit_pairs:
            for ID_digit_pair in ID_digit_pairs:
                if comparand_digit_pair == ID_digit_pair:
                    similar_pairs += 1
        similarity = similar_pairs/len(comparand_digit_pairs)
        #If the percentage of similar digit pairs in the comparand ID matches the desired similarity percentage, add it to the similar ID array
        if similarity >= similarity_percentage:
            similar_IDs.append(comparand_ID)
        row_index += 1
    print("found " + str(len(similar_IDs)) + " similar IDs")
    return similar_IDs

def MatchDataToIDs(data, obj_type, data_type = "null"):
    '''
    Use data to find correlated IDs, and if avaliable use the specified data type to speed up the search. Do this by:
        First, preparing the save file to be read.
        Second, if the data type was specified, identify its correlated key column, otherwise iterate though every key column 
        Third, iterate through the rows of a given key column and add rows with a match to an array
        Fourth, iterate through the array of matching rows and add their correlated IDs to an array
        Fifth, return the ID array
    '''
    
    #Get the object type, get the save file for the object type, and then prepare the save file to be read 
    save_file_PATH = GetSaveFile(obj_type)
    save_file = load_workbook(save_file_PATH)
    output_data = save_file.active
    
    data_column = "null"
    keys = GetObjKeys("null", "null", obj_type)
    key_columns = GetKeyColumns(obj_type)
    
    #If the data type was specified, identify which key column correlates to the data type
    if data_type != "null":
        key_index = 0
        for key in keys:
            key = str(key)
            is_data_type = key.find(data_type)
            if is_data_type != -1:
                data_type_index = key_index
                break
            else:
                key_index += 1
        data_column = key_columns[data_type_index]

        #Iterate through the data type's rows and add each matching row to an array
        row_index = 2
        data_match_rows = []
        row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
        while row_index != row_max:
            cell = data_column + str(row_index)
            cell_value = output_data[cell].value
            if str(cell_value) != "None":
                is_data = cell_value.find(data)
                if is_data != -1:
                    data_match_rows.append(row_index) 
            row_index += 1
        
        #Iterate through the array of matching rows and identify the IDs of the objects that contain the matching data
        ID_matches = []
        for row in data_match_rows:
            cell = "A" + str(row)
            ID = output_data[cell].value
            ID_matches.append(ID)
            
    #Iterate through all key columns of an object type if a data type is not specified 
    else:
        data_match_rows = []
        row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
        #Iterate through the rows of every key column and add any matching rows to an array
        for column in key_columns:
            print("searching column: " + column)
            row_index = 2
            while row_index != row_max:
                cell = column + str(row_index)
                cell_value = output_data[cell].value
                if str(cell_value) != "None":
                    is_data = cell_value.find(data)
                    if is_data != -1:
                        data_match_rows.append(row_index) 
                row_index += 1
                
        #Iterate through the array of matching rows and identify the IDs of the objects that contain the matching data
        ID_matches = []
        for row in data_match_rows:
            cell = "A" + str(row)
            ID = output_data[cell].value
            ID_matches.append(ID)
    return ID_matches
        
    #Find the ID's row in the ID column and then use it to retrieve the object's row   
    row_index = 2
    data_cell = "null"
    data_cell_found = False
    row_max = 1048577 #Only 1,048,576 rows in a spreadsheet
    '''
    while data_cell_found == False and row_index != row_max:
        cell = "A" + str(row_index)
        if output_data[cell].value == ID:
            ID_cell_found = True
            ID_cell = cell
            obj_row_index = row_index 
            break 
        row_index += 1
    if ID_cell == "null":
        print("Error: object does not exist or object's ID is corrupted")
        return "null"
    '''

def CreateObjType(obj_type, obj_type_keys): 
    '''This function creates new objects automatically. Since this function updates existing code, it's work is first saved in the update file. To enact the update it must be imported to the main file via ImportUpdate(). Should you wish to remove the update, you can invoke ImportBackup(), which will revert both the update and main file back to their previous renditions.'''
    
    #Export the latest version of the main script to the extension file 
    ExportMainToUpdate()
    
    single_indent = "    "
    double_indent = "        "
    triple_indent = "            "
    with open(update_script, "r") as update:
        update_txt = update.read() #Get all of the text in the update file as a string
        
        #Create a class declaration for this object and iniate the desired keys 
        classes_start = update_txt.find("#Classes") + len("#Classes")
        classes_end = update_txt.find("#Classes", classes_start)
        classes = update_txt[classes_start:classes_end] #The position where class declarations occur 
        classes = classes + "class " + obj_type + ":" + newline + single_indent + "def __init__(self):" + newline + double_indent #Declaring the class
        classes = classes + "self." + "obj_name" + " = []" + newline + double_indent#Add the object name key to all classes
        for key in obj_type_keys: #Iterating through the desired keys and iniating them 
            if key == obj_type_keys[len(obj_type_keys) - 1]:
                classes = classes + "self." + key + " = []" + newline #prevents the classes comment from being formatted incorrectly 
            else:
                classes = classes + "self." + key + " = []" + newline + double_indent
        update_first_half = update_txt[0:classes_start]
        update_second_half = update_txt[classes_end:len(update_txt)]
        update_txt = update_first_half + classes + update_second_half #Saving the changes 
        
        #Add the object's type tag to the Generate ID function
        GenerateID_start = update_txt.find("def GenerateID(obj):")
        GenerateID_end = update_txt.find("#END", GenerateID_start)
        GenerateID = update_txt[GenerateID_start:GenerateID_end]
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = GenerateID.find("elif obj_type ==", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        last_elif_start = last_elif_pos
        last_type_tag_pos = GenerateID.find("type_tag = ", last_elif_start) + len("type_tag = ")
        last_type_tag = GenerateID[last_type_tag_pos:last_type_tag_pos + 2]#Get the two digit type type
        last_type_tag = int(last_type_tag)
        type_tag = str(last_type_tag + 1)
        end_indicator_length = len("#elif_end")
        newline_length = len(newline)
        last_elif_end = GenerateID.find("#elif_end", last_elif_start) + end_indicator_length + newline_length
        GenerateID = GenerateID[0:last_elif_end] + newline + single_indent + "elif obj_type == " + '"' + obj_type + '"' + ":" + newline + double_indent + "type_tag = " + type_tag + "   #elif_end" + GenerateID[last_elif_end:len(GenerateID)]#Add the new type tag to the elif expression      
        update_first_half = update_txt[0:GenerateID_start]
        update_second_half = update_txt[GenerateID_end:len(update_txt)]
        update_txt = update_first_half + GenerateID + update_second_half#Save the changes made to GenerateID()
        
        #Add the object's type tag to the GetObjType() function 
        GetObjType_start = update_txt.find("def GetObjType(obj = ")
        GetObjType_end = update_txt.find("#END", GetObjType_start)
        GetObjType = update_txt[GetObjType_start:GetObjType_end]
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = GetObjType.find("elif type_tag == ", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        last_elif_start = last_elif_pos
        last_elif_end = GetObjType.find("#elif_end", last_elif_start) + len("#elif_end") + len(newline)
        GetObjType = GetObjType[0:last_elif_end] + newline + double_indent + "elif type_tag == " + '"' + type_tag + '"' + ":" + newline + triple_indent + "obj_type = " + '"' + obj_type + '"' + "#elif_end" + GetObjType[last_elif_end:len(GetObjType)]#Add the new type to the elif expression
        update_first_half = update_txt[0:GetObjType_start]
        update_second_half = update_txt[GetObjType_end:len(update_txt)]
        update_txt = update_first_half + GetObjType + update_second_half#Save the changes made to GetObjType()
        
        #Add the object's type to the CreateEmptyObj() function
        CreateEmptyObj_start = update_txt.find("def CreateEmptyObj(obj_type)")
        CreateEmptyObj_end = update_txt.find("#END", CreateEmptyObj_start)
        CreateEmptyObj = update_txt[CreateEmptyObj_start:CreateEmptyObj_end]
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = CreateEmptyObj.find("elif obj_type == ", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        last_elif_start = last_elif_pos
        last_elif_end = CreateEmptyObj.find("()", last_elif_start) + len("()")
        CreateEmptyObj = CreateEmptyObj[0:last_elif_end] + newline + single_indent + "elif obj_type == " + '"' + obj_type + '"' + ":" + newline + double_indent + "obj = " + obj_type + "()" + CreateEmptyObj[last_elif_end:len(CreateEmptyObj)]
        update_first_half = update_txt[0:CreateEmptyObj_start]
        update_second_half = update_txt[CreateEmptyObj_end:len(update_txt)]
        update_txt = update_first_half + CreateEmptyObj + update_second_half#Save the changes made to GetObjType()
        
        #Create save file
        CreateSaveFile(obj_type, obj_type_keys)
        
        #Overwrite previous update file with the new update 
        with open(update_script, "w") as update: 
            update.write(update_txt)
        print("Created the " + obj_type + " object type. Invoke ImportUpdate() to save its addition.")

def DeleteObjType(obj_type):
    '''
    This function deletes an object type by:
        1. Deleting its class declaration
        2. Deleting it from the CreateEmptyObj() function
        3. Deleting it from the GenerateID() function, and then revising the type tags that follow it to reflect its abscence 
        4. Deleting it from the GetObjType() function, and then revising the type tags that follow it to reflect its abscence
    All updates made to the source code are saved in the update file. Therefore, for the deletion to be enacted, the ImportUpdate() function must be invoked. 
    '''
    
    if obj_type == "Doc":
        print("This object type cannot be deleted, as it is essential for the program to function as intended")
        return
    if CreateEmptyObj(obj_type) == "null":
        print("error: object cannot be deleted, as object does not exist")
        return
        
    ExportMainToUpdate()   
    single_indent = "    "
    double_indent = "        "
    triple_indent = "            "
    with open(update_script, "r") as update:
        update_txt = update.read() #Get all of the text in the update file as a string    
        
        #Delete the class declaration for the object type
        classes_start = update_txt.find("#Classes") + len("#Classes")
        classes_end = update_txt.find("#Classes", classes_start)
        classes = update_txt[classes_start:classes_end] #The position where class declarations occur 
        obj_type_start = classes.find("class " + obj_type)
        obj_keys = GetObjKeys("null", "null", obj_type)
        last_key = obj_keys[len(obj_keys) - 1]
        obj_type_end = classes.find(last_key) + len(last_key) + len(" = []") + len(newline)
        classes = classes[0:obj_type_start] + classes[obj_type_end:len(classes)] #Delete the object type's class declaration by omission 
        update_first_half = update_txt[0:classes_start]
        update_second_half = update_txt[classes_end:len(update_txt)]
        update_txt = update_first_half + classes + update_second_half #Save the deletion 
        
        #Delete the object type from the CreateEmptyObj() function
        CreateEmptyObj_start = update_txt.find("def CreateEmptyObj(obj_type)")
        CreateEmptyObj_end = update_txt.find("#END", CreateEmptyObj_start)
        CreateEmptyObj_ = update_txt[CreateEmptyObj_start:CreateEmptyObj_end]
        obj_type_end = CreateEmptyObj_.find(obj_type + "()") + len(obj_type) + len("()") + len(newline)
        last_elif_pos = 0
        last_elif_located = False 
        while last_elif_located == False: #Iterate through the elif expression until the last type tag in the expression is found
            else_if_pos = CreateEmptyObj_[0:obj_type_end].find(single_indent + "elif obj_type == ", last_elif_pos)
            if else_if_pos != -1:
                last_elif_pos = else_if_pos + 1
            else:
                last_elif_located = True
                break
        obj_type_start = last_elif_pos - 1
        CreateEmptyObj_ = CreateEmptyObj_[0:obj_type_start] + CreateEmptyObj_[obj_type_end:len(CreateEmptyObj_)]
        update_first_half = update_txt[0:CreateEmptyObj_start]
        update_second_half = update_txt[CreateEmptyObj_end:len(update_txt)]
        update_txt = update_first_half + CreateEmptyObj_ + update_second_half
        
        #Delete the object type from GenerateID()
        GenerateID_start = update_txt.find("def GenerateID(obj):")
        GenerateID_end = update_txt.find("#END", GenerateID_start)
        GenerateID = update_txt[GenerateID_start:GenerateID_end]
        obj_type_start = GenerateID.find(single_indent + "elif obj_type == " + '"' + obj_type + '"')
        obj_type_end = GenerateID.find("#elif_end", obj_type_start) + len("#elif_end") + len(newline) + len(newline)
        obj_type_pos = GenerateID[obj_type_start:obj_type_end]
        type_tag_start = obj_type_pos.find("type_tag = ") + len("type_tag = ")
        type_tag_end = type_tag_start + 2
        type_tag = obj_type_pos[type_tag_start:type_tag_end]
        type_tag = int(type_tag)
        type_tags_revised = False
        GenerateID = GenerateID[0:obj_type_start] + GenerateID[obj_type_end:len(GenerateID)]
        type_tag_index = 1
        elif_end = GenerateID.find("#Ensures that null is returned when the object type is invalid")
        elif_pos = GenerateID[0:elif_end]
        while type_tags_revised == False: #Iterate through the type tags after the the object type's tag and subtract one from them to reflect the one less type tag
            type_tag_to_revise_start = elif_pos.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 2
            type_tag_to_revise_end = elif_pos.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 3
            type_tag_to_revise = elif_pos[type_tag_to_revise_start] + elif_pos[type_tag_to_revise_end]
            try:
                type_tag_to_revise = int(type_tag_to_revise) #If the integer is invalid then there are no more type tags to revise. Therefore, break the loop
                GenerateID = GenerateID.replace(str(type_tag_to_revise), str(type_tag_to_revise - 1))
                type_tag_index += 1
            except:
                type_tags_revised = True
                break      
        update_first_half = update_txt[0:GenerateID_start]
        update_second_half = update_txt[GenerateID_end:len(update_txt)]
        update_txt = update_first_half + GenerateID + update_second_half

        #Delete the object type from the GetObjType() function
        GetObjType_start = update_txt.find("def GetObjType(obj = ")
        GetObjType_end = update_txt.find("#END", GetObjType_start)
        GetObjType = update_txt[GetObjType_start:GetObjType_end]
        obj_type_start = GetObjType.find(single_indent + "elif type_tag == " + '"' + str(type_tag) + '"') - len(single_indent)
        if obj_type_start == (-1 - len(single_indent)):
            obj_type_start = GetObjType.find(single_indent + "elif type_tag == " + '"' + "0" + str(type_tag) + '"') - len(single_indent)
        obj_type_end = GetObjType.find("obj_type = " + '"' + obj_type + '"' + "#elif_end") + len("obj_type = " + '"' + obj_type + '"' + "#elif_end" + newline) 
        GetObjType = GetObjType[0:obj_type_start] + GetObjType[obj_type_end:len(GetObjType)]   
        type_tags_revised = False
        type_tag_index = 1
        while type_tags_revised == False: #Iterate through the type tags after the the object type's tag and subtract one from them to reflect the one less type tag
            type_tag_to_revise_start = GetObjType.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 2
            type_tag_to_revise_end = GetObjType.find(str(type_tag + type_tag_index)) - len(str(type_tag + type_tag_index)) + 3
            type_tag_to_revise = GetObjType[type_tag_to_revise_start] + GetObjType[type_tag_to_revise_end]
            try:
                type_tag_to_revise = int(type_tag_to_revise) #If the integer is invalid then there are no more type tags to revise. Therefore, break the loop
                if type_tag_to_revise - 1 < 10: #If the type tag is less than 10, add a leading zero
                    GetObjType = GetObjType.replace(str(type_tag_to_revise), "0" + str(type_tag_to_revise - 1))
                else:
                    GetObjType = GetObjType.replace(str(type_tag_to_revise), str(type_tag_to_revise - 1))
                type_tag_index += 1
            except:
                type_tags_revised = True
                break    
        update_first_half = update_txt[0:GetObjType_start]
        update_second_half = update_txt[GetObjType_end:len(update_txt)]
        update_txt = update_first_half + GetObjType + update_second_half
        
        #Delete save file
        save_file = GetSaveFile(obj_type)
        os.remove(save_file)

        #Overwrite previous update file with the new update 
        with open(update_script, "w") as update: 
            update.write(update_txt)
        print("Deleted the " + obj_type + " object type. Invoke ImportUpdate() to save its deletion.")
     
def ExportMainToBackup():
    copy(main_script, backup_script)
    print("exported main to backup")
     
def ImportBackup():
    copy(backup_script, main_script)
    copy(backup_script, update_script)
    print("imported backup to main")

def ImportUpdate():
    ExportMainToBackup()
    copy(update_script, main_script)
    print("imported update to main")
    
def ExportMainToUpdate():
    copy(main_script, update_script)
    print("exported main to update")

def GetObjTypeNames():
    obj_type_index = 0
    obj_type_names = []
    get_obj_type_names = True
    while get_obj_type_names == True:
        if obj_type_index < 10:
            obj_type_name = GetObjType("null", "0" + str(obj_type_index))
            if obj_type_name == "null":
                break
            else:
                obj_type_names.append(obj_type_name)
        else:
            obj_type_name = GetObjType("null", str(obj_type_index))
            if obj_type_name == "null":
                break
            else:
                obj_type_names.append(obj_type_name)
        obj_type_index +=1
    return obj_type_names

#functions

#def AddCharacterToInputValues(), adds a character to the list of characters allowed to be used for value input by adding the character to the TxtToDigit and DigitToNum functions

#def DeleteObjKey

#def GetDocFromDocObj

#def AddObjKey 


#TEST CODE GOES HERE!!!


'''
DO GUI NEXT!!!
class Compound_Conditional_Instruction:
    def __init__(self):
        self.obj_name = []
        self.conditional_instructions = []
        

class Conditional_Instruction:   
    def __init__(self):
        self.obj_name = []
        self.condition_type = []
        self.condition = []
        self.instruction = []
        self.place_holder = []
    def CheckConditions(self, parameters):
        for condition in self.condition_type:
            #every parameter is equal to the condition
            if condition == "==":
                condition_met = True
                for parameter in parameters:
                    if parameter == self.condition[0]:
                        condition_met = True
                    else:
                        condition_met = False
                        break
            #every parameter is NOT equal to the condition
            elif condition == "!=":
                condition_met = True
                for parameter in parameters:
                    if parameter != self.condition[0]:
                        condition_met = True
                    else:
                        condition_met = False
                        break
        return condition_met
    def ExportValueToPlaceHolder(self, value, doc_path):
        FormatKeys(doc_path)
        doc = Document(doc_path)
        key = key_start_indicator + self.place_holder[0] + key_end_indicator
        for paragraph in doc.paragraphs:
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
        doc.save(doc_path)
    def EvaluateExpression(self, parameters, doc_path):
        if self.CheckConditions(parameters) == True:
            self.ExportValueToPlaceHolder(self.instruction[0], doc_path)
        else:
            FormatKeys(doc_path)
            doc = Document(doc_path)
            key = key_start_indicator + self.place_holder[0] + key_end_indicator
            for paragraph in doc.paragraphs:
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, "")
            doc.save(doc_path)
    

doc_path = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "templates\\Doc.docx")
sublang = Conditional_Instruction()
sublang.obj_name = ["sublang"]
sublang.condition_type = ["=="]
sublang.condition = ["employer"]
sublang.instruction = ["Give me employer records"]
sublang.place_holder = ["poop"]


poop = Entity()
poop.entity_type = ["employer"]



sublang.EvaluateExpression([poop.entity_type[0]], doc_path)

'''




















'''Add deletion of excel file to deletion code but otherwise GOOD'''

'''
doc_path = MAIN_FILE_PATH.replace(MAIN_FILE_NAME, "templates\\Doc.docx")
doc = Document(doc_path)




workbook = load_workbook("C:\\Development\\0.0 RawPrograms\\DocGen\\data\\lawyer_data.xlsx")
sheet = workbook.active



sheet["A2"] = "Hello"
sheet["A3"] = "shit"
sheet["B1"] = "world!"

workbook.save("C:\\Development\\0.0 RawPrograms\\DocGen\\data\\lawyer_data.xlsx")

print(sheet["A2"].value)

'''


#TEST CODE GOES HERE!!!

   
#Notes
'''
All Documents have to be a docx document
'''
#Notes
   
   
   
   
   
        #def callback(sender, app_data, user_data):, saved for reference
        #print(f"user_data is: {user_data}"), saved for reference
        #print(f"sender is: {sender}"), saved for reference
        #print(f"app_data is: {app_data}"), saved for reference
        


#GUI Color Palette
latte = [210,191,168]
light_cream = [255,253,245]
heavy_custard = [245, 240, 208]
deep_ocean_blue = [95,143,155]
seaglass = [169,209,217]
black = [0, 0, 0]
white = [255,255,255]
dark_gray = [69, 69, 65]
gray = [170, 170, 170]
light_gray = [230, 231, 227]
dark_blue = [86, 104, 172]
blue = [120, 136, 218]
light_blue = [160, 205, 240]
dark_green = [33, 78, 63]
green = [96, 153, 123]
sea_foam = [140, 228, 168]
coral = [244, 193, 193]
brown = [185, 143, 123]
pink = [255, 163, 192]
yellow = [235, 203, 113]
apricot = [252, 223, 179]
white_apricot = [255, 255, 235]



primary_window = "DefaultWindow" #The primary window/fullscreen window
def ResizeRows(parent_table):
    default_widget_height = 31.775
    rows = dpg.get_item_children(parent_table, slot=1)
    row_height = default_widget_height
    for row in rows:
        dpg.configure_item(row, height=row_height)  
    return 


def CreateTable(table_tag, window, column_amount=1, row_amount=1):
    
    #Creates a table though assumes only one table per window
    #the "TablePos" group moves the table, and each cell of the 
    #the table can be accessed via the following format: "Row1Col2Content"
    
    with dpg.group(tag="TablePos", pos=[0, 0], parent=window):    
        with dpg.table(tag=table_tag, header_row=False, width=0, borders_outerH=True, borders_outerV=True, borders_innerH=True, borders_innerV=True, policy=dpg.mvTable_SizingStretchProp, precise_widths=True, scrollY=True):
            dpg.bind_item_theme(table_tag, table_background_theme)
            AppendCells(table_tag, column_amount=column_amount, row_amount=row_amount)
            
def AppendCells(parent_table, column_amount=1, row_amount=1):
    #Estabilishes a default width weight and height
    column_width_weight = 1.0
    row_height = 1
    
    #Determine Column Index
    column_index=1
    while dpg.does_item_exist("Col" + str(column_index)) == True:
        column_index+=1 
    
    #Determine Row Index
    row_index=1
    while dpg.does_item_exist("Row" + str(row_index)) == True:
        row_index+=1 

    #Populate the columns 
    if column_amount != 0:
        while column_index <= column_amount:
            dpg.add_table_column(tag="Col" + str(column_index), init_width_or_weight=column_width_weight, parent=parent_table)
            column_index += 1
        if column_index != 1:
            column_amount = column_amount + column_index - 1
    else: 
        column_amount = column_index - 1

    #Populate the rows  
    if row_amount != 0:
        if row_index != 1:
            row_amount = row_amount + row_index - 1
        while row_index <= row_amount:
            with dpg.table_row(tag="Row" + str(row_index), height=row_height, parent=parent_table):
                column_index = 1
                while column_index <= column_amount:
                    with dpg.group(tag="Row" + str(row_index) + "Col" + str(column_index) + "Content"): 
                        dpg.add_spacer(width=0, height=0)
                    column_index+=1
                row_index += 1
    ResizeRows(parent_table)
    
    row_index = 1
    while row_index <= row_amount:
        cells = dpg.get_item_children("Row" + str(row_index), slot=1)
        for cell in cells:
            cell_tag = dpg.get_item_alias(cell)
            cell_col = int(cell_tag[(cell_tag.find("Col") + len("Col")):cell_tag.find("Content")])
            if cell_col > column_amount:
                dpg.delete_item(cell_tag)
        row_index += 1
                    
def InsertRow(parent_table, insert_pos=1, row_amount=1):
    inserted_row_amount = row_amount
    
    #Determine Row Amount
    row_index=1
    while dpg.does_item_exist("Row" + str(row_index)) == True:
        row_index+=1 
    row_amount = row_index - 1
    
    #Determine Column Amount
    col_index=1
    while dpg.does_item_exist("Col" + str(col_index)) == True:
        col_index+=1 
    col_amount = col_index - 1

    #Reassign aliases to pre-existing rows 
    row_index = row_amount 
    while row_index >= 1:
        if row_index >= insert_pos:
            dpg.set_item_alias("Row" + str(row_index), "Row" + str(row_index + inserted_row_amount))
            col_index = col_amount
            while col_index >= 1:
                dpg.set_item_alias("Row" + str(row_index) + "Col" + str(col_index) + "Content", "Row" + str(row_index + inserted_row_amount) + "Col" + str(col_index) + "Content")
                col_index -= 1
        else:
            pass
        row_index -= 1

    #Insert the new rows  
    row_height = 1
    row_index = insert_pos
    while row_index <= (insert_pos + inserted_row_amount - 1):
        if dpg.does_item_exist("Row" + str(row_index)):
            dpg.remove_alias("Row" + str(row_index))
        with dpg.table_row(tag="Row" + str(row_index), height=row_height, parent=parent_table, before="Row" + str(insert_pos + inserted_row_amount)):
            col_index = 1
            while col_index <= col_amount:
                if dpg.does_item_exist("Row" + str(row_index) + "Col" + str(col_index) + "Content"):
                    dpg.remove_alias("Row" + str(row_index) + "Col" + str(col_index) + "Content")
                with dpg.group(tag="Row" + str(row_index) + "Col" + str(col_index) + "Content"): 
                    dpg.add_spacer(width=0, height=0)
                col_index+=1
            row_index += 1
    ResizeRows(parent_table)
    rows = dpg.get_item_children(parent_table, slot=1)

    
    

    

 




    
    
def DefaultWindow():
    with dpg.window(tag="DefaultWindow") as default_window:
        dpg.add_button(
        label="New Obj",
        width=185, 
        height=31.775, 
        parent="DefaultWindow",
        callback=NewObj_Window,
        pos=(35, 50))
        
#The code below is basically done but needs to be rewritten for table centric structure 
def CreateObjWindow_defunct(calc_offset=True, saved_input=["null"]):
    global primary_window
    offset=0
    last_value_IDs=[] 

    if primary_window == "SaveObj" and calc_offset==True:
        MainWindow()
        dpg.set_primary_window("PrimaryWindow", True)
        offset = (dpg.get_viewport_width())/4
        max_viewport_width = dpg.get_viewport_width()
        max_viewport_height = dpg.get_viewport_height()
        if max_viewport_width < max_viewport_height:
            max_viewport_width = max_viewport_width * 2 
        elif max_viewport_width < 800: #only effective for 92% of monitors 
            max_viewport_width = max_viewport_width * 2
        elif max_viewport_height < 600: #only effective for 92% of monitors 
            max_viewport_width = max_viewport_width * 2
            
        dpg.set_primary_window("PrimaryWindow", False)
        dpg.delete_item("PrimaryWindow")
    elif primary_window == "SaveObj" and calc_offset == False:
        max_viewport_width = dpg.get_viewport_width() * 2
    
    #Themes
    with dpg.theme() as add_subfield_button_theme:
        with dpg.theme_component(dpg.mvAll):
            dpg.add_theme_color(dpg.mvThemeCol_Button, (sea_foam[0], sea_foam[1], sea_foam[2]), category=dpg.mvThemeCat_Core) 
    with dpg.theme() as delete_subfield_button_theme:
        with dpg.theme_component(dpg.mvAll):
            dpg.add_theme_color(dpg.mvThemeCol_Button, (coral[0], coral[1], coral[2]), category=dpg.mvThemeCat_Core)        
    with dpg.theme() as sub_text_button_theme:
        with dpg.theme_component(dpg.mvAll):
            dpg.add_theme_color(dpg.mvThemeCol_Button, (white_apricot[0], white_apricot[1], white_apricot[2]), category=dpg.mvThemeCat_Core)        
    with dpg.theme() as text_button_theme:
        with dpg.theme_component(dpg.mvAll):
            dpg.add_theme_color(dpg.mvThemeCol_Button, (apricot[0], apricot[1], apricot[2]), category=dpg.mvThemeCat_Core)
    with dpg.theme() as menu_bar_drop_down_theme:
        with dpg.theme_component(dpg.mvAll):
            dpg.add_theme_color(dpg.mvThemeCol_Button, (creamer[0], creamer[1], creamer[2]), category=dpg.mvThemeCat_Core) 
            dpg.add_theme_color(dpg.mvThemeCol_FrameBg, (creamer[0], creamer[1], creamer[2]), category=dpg.mvThemeCat_Core) #drop down menu button backround
            dpg.add_theme_color(dpg.mvThemeCol_Header, (creamer[0], creamer[1], creamer[2]), category=dpg.mvThemeCat_Core) #drop down menu item color
            dpg.add_theme_color(dpg.mvThemeCol_HeaderHovered, (latte[0], latte[1], latte[2]), category=dpg.mvThemeCat_Core) #drop down menu item color
            dpg.add_theme_color(dpg.mvThemeCol_HeaderActive, (latte[0], latte[1], latte[2]), category=dpg.mvThemeCat_Core) #drop down menu item color
            dpg.add_theme_color(dpg.mvThemeCol_PopupBg , (creamer[0], creamer[1], creamer[2]), category=dpg.mvThemeCat_Core) #drop down menu 

            
    def GetUserInput():
        '''This function retrieves user input and stores it in an array in the event of a window resize or respawn'''
        
        window_children = dpg.get_item_children("SaveObj").get(1)
        child_index = 0
        for child in window_children:
            window_children[child_index] = dpg.get_item_alias(child)
            child_index += 1
        value_children = []
        if "GetKeyValues" in window_children:
            value_children = dpg.get_item_children("GetKeyValues").get(1)
            
        if len(value_children) >= 1:
            value_child_index = 0
            for value_child in value_children:
                value_children[value_child_index] = dpg.get_item_alias(value_child)
                value_child_index += 1
            
            value_child_index = 0
            for value_child in value_children:
                value_children[value_child_index] = [value_child, dpg.get_value(value_child)]
                value_child_index += 1
                    
            value_child_index = 0
            flagged_children = []
            while value_child_index < len(value_children):
                value_child = value_children[value_child_index]
                value_child_index += 1
                if "EnterKey" not in value_child[0] or value_child[1] == None:
                    flagged_children.append(value_child)
                    
            for flagged_child in flagged_children:
                value_children.remove(flagged_child)

            ChooseObjType = []
            for window_child in window_children:
                if window_child == "ChooseObjType":
                    ChooseObjType = [window_child, dpg.get_value(window_child)]
            value_children.append(ChooseObjType)
        else:
            value_children = []
            ChooseObjType = []
            for window_child in window_children:
                if window_child == "ChooseObjType":
                    ChooseObjType = [window_child, dpg.get_value(window_child)]
            value_children.append(ChooseObjType)
        return value_children 
    
    #If the user tries to open this window when it is already open, simply respawn the same window
    if dpg.does_item_exist("SaveObj") == True:
        saved_input = GetUserInput()
        dpg.delete_item("SaveObj")
        CreateObjWindow(calc_offset=True, saved_input=saved_input)    
        return 
   
    def AutoSizeWindow(): 
        nonlocal offset
        nonlocal max_viewport_width
        if primary_window == "SaveObj":
            if offset != 0:
                if dpg.get_viewport_width() != max_viewport_width:
                    saved_input = GetUserInput()
                    dpg.delete_item("SaveObj")
                    CreateObjWindow(calc_offset=False, saved_input=saved_input)
            elif offset == 0:
                width_to_max_width = dpg.get_viewport_width()/max_viewport_width
                max_width_to_width = max_viewport_width/dpg.get_viewport_width()
                if (width_to_max_width >= .90 and width_to_max_width <= 1.0) or (max_width_to_width >= .90 and max_width_to_width <= 1.0):
                    saved_input = GetUserInput()
                    dpg.delete_item("SaveObj")
                    CreateObjWindow(calc_offset=True, saved_input=saved_input)
                    
    def GenObjValueFields(sender, app_data, user_data):
        '''
        This function receives the object type specified by the user and 
        then generates value entry fields pertaining to that object type
        '''
        
        obj_type = app_data
        
        #Get the keys/attributes of the object type
        obj_type_keys = GetObjKeys("null", "null", obj_type)
        
        #dynamically creates the value entry fields
        input_text_x_position = 225
        input_text_y_position = 125
        
        #if the value entry fields already exist, due to a pre-existing object type choice, delete them and replace them with the value entry fields pertaining to the newly selected object type
        if dpg.does_item_exist("GetKeyValues") == True:
            dpg.delete_item("GetKeyValues")  
        
        with dpg.group(tag="GetKeyValues", parent="SaveObj") as GetKeyValues:    
            
            #create a value entry field for each of the object type's attributes/key
            clean_keys = []
            key_ID = 1
            for key in obj_type_keys:
                #if the key contains more than one word, clean up the programmer portrayal of the sequence into a more pallatable one fit for a user
                if "_" in key: #if an underscore is present, there is more than one word
                    clean_key = [] #an array of the neatly seperated words
                    last_seperator_pos = 0
                    seperator_index = 0
                    #fetch the seperate words from the key and then store them in the clean key array
                    while seperator_index <= key.count("_"):
                        if seperator_index == key.count("_"):
                            seperator_pos = len(key)
                            last_seperator_pos += 1
                        else:
                            seperator_pos = key.find("_", last_seperator_pos)
                        clean_key.append(key[last_seperator_pos:seperator_pos])
                        last_seperator_pos = seperator_pos
                        seperator_index += 1
                    #iterate through the array of seperated words, and clean them even further by applying desirable formatting 
                    key_part_index = 0
                    for key_part in clean_key:
                        key_part = key_part[0].upper() + key_part[1:len(key_part)]
                        clean_key[key_part_index] = key_part
                        key_part_index += 1
                    key = ""
                    for key_part in clean_key:
                        if key_part != clean_key[len(clean_key) - 1]:
                            key = key + key_part + " "
                        else:
                            key = key + key_part + ":"
                    clean_keys.append(key)
                else:
                    #apply desirable formatting to the single worded key/attribute
                    key = key[0].upper() + key[1:len(key)] + ":"
                    clean_keys.append(key)

                #create the text preceding the value entry field that identifies what the field pertains to 
                dpg.add_button(
                label=key,
                width=185, 
                height=31.775,
                tag="EnterKeyTxt#" + str(key_ID), 
                parent="GetKeyValues",
                pos=(offset + 35, input_text_y_position))
                
                dpg.bind_item_theme("EnterKeyTxt#" + str(key_ID), text_button_theme)
                
                #create the value entry field
                dpg.add_input_text(
                tag="EnterKey#" + str(key_ID),
                parent="GetKeyValues", 
                width=495.5, 
                height=600, 
                on_enter=True,
                callback=JumpToNextFieldOnEnter,
                pos=(offset + input_text_x_position, input_text_y_position))
                
                dpg.add_button(
                tag="AddKey#" + str(key_ID),
                label="Add Value",
                width=101.5, 
                height=31.775, 
                parent="GetKeyValues", 
                pos=(offset + 725, input_text_y_position),
                user_data=[key, key_ID, input_text_x_position, input_text_y_position],
                callback=CreateFieldForNewArrayMember)

                dpg.bind_item_theme("AddKey#" + str(key_ID), add_subfield_button_theme)
                input_text_y_position += 50
                key_ID += 1
            
            dpg.add_button(
            label="Save Object", 
            tag="SaveObjButton",
            parent="GetKeyValues", 
            pos=(offset + 35, input_text_y_position), 
            width=795.5, 
            height=50,
            user_data=[obj_type, obj_type_keys],
            callback=SaveObjFrUser)
            
            
            dpg.add_text("", tag="BottomBuffer", parent="GetKeyValues", pos=(350, input_text_y_position + 75))
    
            sub_field_input_exists = False
            if saved_input == None:
                pass
            else: 
                for user_input in saved_input:
                    if user_input[0] == "ChooseObjType":
                        pass
                    elif "!" in user_input[0]:
                        sub_field_input_exists = True
                        pass
                    else:
                        dpg.set_value(user_input[0], user_input[1])
                if sub_field_input_exists == True:
                    for user_input in saved_input:
                        if "!" in user_input[0]:
                            key_ID_start = user_input[0].find("#") + 1
                            key_ID_end = user_input[0].find("!")
                            key_ID = user_input[0][key_ID_start:key_ID_end]
                            key = dpg.get_item_label("EnterKeyTxt#" + key_ID)
                            input_text_x_position = dpg.get_item_pos("EnterKey#" + key_ID)[0]
                            if primary_window == "SaveObj":
                                input_text_x_position = input_text_x_position - offset
                            input_text_y_position = dpg.get_item_pos("EnterKey#" + key_ID)[1]
                            CreateFieldForNewArrayMember("null", "null", [key, key_ID, input_text_x_position, input_text_y_position])
                            dpg.set_value(user_input[0], user_input[1])
                    
    def CreateFieldForNewArrayMember(sender, app_data, user_data):
        '''
        Create a new sub input field to accept more than one value for a given key
        '''
        nonlocal last_value_IDs
        #Organize the user data
        key = user_data[0]
        key_ID = int(user_data[1])
        input_text_x_position = user_data[2]
        input_text_y_position = user_data[3]
        
        #Find the next avaliable value ID for a sub input field 
        value_ID = "null"
        real_value_ID = "null"
        try_value_ID = 1
        while value_ID == "null":
            try:
                dpg.get_item_type("EnterKey#" + str(key_ID) + "!" + str(try_value_ID))
                try_value_ID += 1
                #add to last value ID
            except:
                last_value_ID = "null"
                value_ID_test = try_value_ID
                if len(last_value_IDs) > 0:
                    value_index = 0
                    for last_value in last_value_IDs:
                        if last_value[0] == key_ID:
                            last_value_ID = last_value[1]
                            break
                        value_index += 1
                    if last_value_ID == "null":
                        last_value_ID = 0
                else:
                    last_value_ID = 0

                if value_ID_test > last_value_ID:
                    value_ID = value_ID_test
                    last_value_ID = value_ID + 1
                    
                    try:
                        last_value_IDs[value_index][1] = last_value_ID
                    except:
                        last_value_IDs.append([key_ID, last_value_ID])

                    if last_value_ID != 2:
                        last_value_ID_index = last_value_ID
                        while real_value_ID == "null" and last_value_ID_index != 0:
                            try:
                                last_label = dpg.get_item_label("EnterKeyTxt#" + str(key_ID) + "!" + str(last_value_ID_index))
                            except:
                                last_label = "null"
                            if last_label != "null": 
                                break
                            else:
                                last_value_ID_index -= 1
                        if last_value_ID_index == 0:
                            real_value_ID = 2
                        else:
                            real_value_ID_start = last_label.find(": ") + 1
                            real_value_ID_end = len(last_label)
                            real_value_ID = int(last_label[real_value_ID_start:real_value_ID_end])
                            real_value_ID += 1
                    else:
                        real_value_ID = 2
                    break 
                else:
                    try_value_ID += 1

        #Create new sub input field
        input_text_y_position = dpg.get_item_pos("EnterKey#" + str(key_ID))[1]
        input_text_y_position = input_text_y_position + ((real_value_ID - 1) * 50)
        dpg.add_input_text(
        tag="EnterKey#" + str(key_ID) + "!" + str(value_ID),
        parent="GetKeyValues", 
        width=495.5, 
        height=600, 
        on_enter=True,
        callback=JumpToNextFieldOnEnter,
        pos=(offset + input_text_x_position, input_text_y_position))

        dpg.add_button(
        label=key + " " + str(real_value_ID),
        width=185, 
        height=31.775,
        tag="EnterKeyTxt#" + str(key_ID) + "!" + str(value_ID), 
        parent="GetKeyValues",
        pos=(offset + 35, input_text_y_position))
                
        dpg.bind_item_theme("EnterKeyTxt#" + str(key_ID) + "!" + str(value_ID), sub_text_button_theme)
                
        dpg.add_button(
        tag="DelKey#" + str(key_ID) + "!" + str(value_ID),
        label="Del Value",
        width=101.5, 
        height=31.775, 
        parent="GetKeyValues", 
        callback=DeleteSubInputField,
        user_data=["EnterKeyTxt#" + str(key_ID) + "!" + str(value_ID), "EnterKey#" + str(key_ID) + "!" + str(value_ID), "DelKey#" + str(key_ID) + "!" + str(value_ID), key_ID, value_ID],
        pos=(offset + 725, input_text_y_position))
        
        dpg.bind_item_theme("DelKey#" + str(key_ID) + "!" + str(value_ID), delete_subfield_button_theme)
        
        #Shift every item such that their original formatting is not disrupted by the new sub input field 
        GetKeyValues_children = dpg.get_item_children("GetKeyValues").get(1)
        entry_fields = []
        for child in GetKeyValues_children:
            if dpg.get_item_alias(child) == "": 
                #null item ID
                pass
            else:
                entry_fields.append(dpg.get_item_alias(child))
                 
        for entry_field in entry_fields:
            try:
                entry_field_ID_start = entry_field.find("#")
                entry_field_ID_end = entry_field.find("!")
                
                if entry_field_ID_end != -1:
                    entry_field_ID = int(entry_field[entry_field_ID_start + 1:entry_field_ID_end])
                else:
                    entry_field_ID = int(entry_field[entry_field_ID_start + 1:len(entry_field)])
                
                if entry_field_ID < key_ID:
                    pass
                elif entry_field_ID == key_ID:
                    field_value_ID_start =  entry_field.find("!")
                    if field_value_ID_start == -1:
                        pass
                    else:
                        field_value_ID_start += 1
                        field_value_ID_end = len(entry_field)
                        field_value_ID = entry_field[field_value_ID_start:field_value_ID_end]
                        field_value_ID = int(field_value_ID)
                        if field_value_ID > value_ID:
                            original_item_position = dpg.get_item_pos(entry_field)
                            original_item_position_x = original_item_position[0]
                            original_item_position_y = original_item_position[1]
                            dpg.set_item_pos(entry_field, [original_item_position_x, original_item_position_y + 50])   
                else:
                    original_item_position = dpg.get_item_pos(entry_field)
                    original_item_position_x = original_item_position[0]
                    original_item_position_y = original_item_position[1]
                    dpg.set_item_pos(entry_field, [original_item_position_x, original_item_position_y + 50])
            except:
                original_item_position = dpg.get_item_pos(entry_field)
                original_item_position_x = original_item_position[0]
                original_item_position_y = original_item_position[1]
                dpg.set_item_pos(entry_field, [original_item_position_x, original_item_position_y + 50])

    def JumpToNextFieldOnEnter(sender, app_data, user_data):
        '''
        Jump to the next input field upon pressing enter
        '''
        
        #Get the children of Get Key Values Group
        GetKeyValues_children = dpg.get_item_children("GetKeyValues").get(1)  
        
        #Find the focused input field among the children, and do not check irrelevant children
        for child in GetKeyValues_children:
            if "EnterKey#" not in dpg.get_item_alias(child): 
                pass
            else:
                if dpg.is_item_focused(dpg.get_item_alias(child)) == True:
                    focused_item = dpg.get_item_alias(child)
        
        #Proceed to the next input field, and if this is the last field, loop back to the first input field
        key_entry_ID_start = focused_item.find("#")
        key_entry_ID_end = focused_item.find("!")
        if key_entry_ID_end == -1:
            key_entry_ID = int(focused_item[key_entry_ID_start + 1:len(focused_item)])
        else:
            key_entry_ID = int(focused_item[key_entry_ID_start + 1:key_entry_ID_end])
            value_entry_ID = int(focused_item[key_entry_ID_end + 1:len(focused_item)])

        if key_entry_ID_end == -1:
            try:
                dpg.focus_item("EnterKey#" + str(key_entry_ID) + "!" + str(1))
            except:
                try:
                    dpg.focus_item("EnterKey#" + str(key_entry_ID + 1))     
                except:
                    dpg.focus_item("EnterKey#" + str(1))                        
        else: 
            try:
                dpg.focus_item("EnterKey#" + str(key_entry_ID) + "!" + str(value_entry_ID + 1))
            except:
                try:
                    dpg.focus_item("EnterKey#" + str(key_entry_ID + 1))
                except:
                    dpg.focus_item("EnterKey#" + str(1))

    def SaveObjFrUser(sender, app_data, user_data):
        #Organize the user data
        obj_type = user_data[0] 
        obj_keys = user_data[1]
        
        #Get the all of children from the GetKeyValues group
        GetKeyValues_children = dpg.get_item_children("GetKeyValues").get(1)  
        
        #Create an array of only the input fields among the children
        input_fields = []
        for child in GetKeyValues_children:
            if "EnterKey#" not in dpg.get_item_alias(child): 
                pass
            else:
                input_fields.append(dpg.get_item_alias(child))

        #Pair the retrieved input fields with their values
        input_field_values = []
        for input_field in input_fields:
            input_field_values.append([input_field, dpg.get_value(input_field)])

        #Organize the input fields into two catergories, fields and subfields
        sub_fields = []
        fields = []
        for field_value_pair in input_field_values:
            field = field_value_pair[0]
            if "!" in field:
                field_ID_start = field.find("#") + 1
                field_ID_end = field.find("!")
                sub_fields.append([int(field[field_ID_start:field_ID_end]), field_value_pair[1]])
            else:
                field_ID_start = field.find("#") + 1
                fields.append([int(field[field_ID_start:len(field)]), field_value_pair[1]])

        #Pair the values of these fields and subfields with their shared key, then create an object
        obj = CreateEmptyObj(obj_type)
        for field in fields:
            key_value_array = []
            key_value_array.append(field[1])
            field_ID = field[0] - 1
            key_index = field_ID
            for sub_field in sub_fields:
                if sub_field[0] == field_ID + 1:
                    key_value_array.append(sub_field[1])
            setattr(obj, obj_keys[key_index], key_value_array)

        #If a value was left empty by the user, set it to null
        for key in obj_keys:
            if getattr(obj, key)[0] == '':
                setattr(obj, key, ["null"])

        #Check if the object already exists, if so, replace only overriden information, leave everything else untouched
        obj_ID = GenerateID(obj)
        if GetObj(obj_ID) == "null":
            print("poop")
            #SaveObj(obj, new_object=True)
        else:
            obj_predecessor = GetObj(obj_ID)
            for key in obj_keys:
                if getattr(obj, key)[0] == "null":
                    setattr(obj, key, getattr(obj_predecessor, key))
            #SaveObj(obj, new_object=False)
            
        #Communicate to the user that the new or overwritten object has been saved via a brief pop-up
        viewport_width = dpg.get_viewport_width()
        viewport_height = dpg.get_viewport_height()
        with dpg.window(tag="ObjSavedPopUp", label="Object Saved", parent="SaveObj", pos=(viewport_width-250,viewport_height-250), no_collapse=True, width=200, height=200) as obj_saved:
            dpg.bind_item_theme(obj_saved, window_theme)#assigns the window theme to this window 
            
    def DeleteSubInputField(sender, app_data, user_data):
        sub_field_txt_tag = user_data[0]
        sub_field_tag = user_data[1]
        sub_field_del_tag = user_data[2]
        key_ID = int(user_data[3])
        value_ID = int(user_data[4])
        sub_field_label = dpg.get_item_label(sub_field_txt_tag)
        dpg.delete_item(sub_field_txt_tag)
        dpg.delete_item(sub_field_tag)
        dpg.delete_item(sub_field_del_tag)
        GetKeyValues_children = dpg.get_item_children("GetKeyValues").get(1)
        child_index = 0
        
        for child in GetKeyValues_children:
            GetKeyValues_children[child_index] = dpg.get_item_alias(child)
            child_index += 1
        new_value_ID = 1
        for child in GetKeyValues_children:
            if "#" in child:
                if "!" in child:
                    comparand_key_ID_start = child.find("#") + 1
                    comparand_key_ID_end = child.find("!")
                    comparand_key_ID = child[comparand_key_ID_start:comparand_key_ID_end]
                    comparand_key_ID = int(comparand_key_ID)
                    if comparand_key_ID > key_ID:
                        original_item_position = dpg.get_item_pos(child)
                        original_item_position_x = original_item_position[0]
                        original_item_position_y = original_item_position[1]
                        dpg.set_item_pos(child, [original_item_position_x, original_item_position_y - 50])
                    elif comparand_key_ID == key_ID:
                        label = dpg.get_item_label(child)
                        if ":" in label: 
                            comparand_value_ID_start = label.find(": ") + 1
                            comparand_value_ID_end = len(label)
                            comparand_value_ID = label[comparand_value_ID_start:comparand_value_ID_end]
                            comparand_value_ID = int(comparand_value_ID)

                            real_value_ID_start = sub_field_label.find(": ") + 1
                            real_value_ID_end = len(sub_field_label)
                            real_value_ID = sub_field_label[real_value_ID_start:real_value_ID_end]
                            real_value_ID = int(real_value_ID)

                            if comparand_value_ID >= real_value_ID:
                                original_item_position = dpg.get_item_pos(child)
                                original_item_position_x = original_item_position[0]
                                original_item_position_y = original_item_position[1]
                                dpg.set_item_pos(child, [original_item_position_x, original_item_position_y - 50])
                                new_label = label[0:comparand_value_ID_start] + " " + str(real_value_ID + (new_value_ID - 1))
                                dpg.configure_item(child, label=new_label)
                                new_value_ID += 1
                        else:
                            comparand_value_ID_start = comparand_key_ID_end + 1
                            comparand_value_ID_end = len(child)
                            comparand_value_ID = child[comparand_value_ID_start:comparand_value_ID_end]
                            comparand_value_ID = int(comparand_value_ID)
                        
                            if comparand_value_ID > value_ID:
                                original_item_position = dpg.get_item_pos(child)
                                original_item_position_x = original_item_position[0]
                                original_item_position_y = original_item_position[1]
                                dpg.set_item_pos(child, [original_item_position_x, original_item_position_y - 50])

                else:
                    comparand_key_ID_start = child.find("#") + 1
                    comparand_key_ID_end = len(child)
                    comparand_key_ID = child[comparand_key_ID_start:comparand_key_ID_end]
                    comparand_key_ID = int(comparand_key_ID)
                    if comparand_key_ID > key_ID:
                        original_item_position = dpg.get_item_pos(child)
                        original_item_position_x = original_item_position[0]
                        original_item_position_y = original_item_position[1]
                        dpg.set_item_pos(child, [original_item_position_x, original_item_position_y - 50])    
            else:
                original_item_position = dpg.get_item_pos(child)
                original_item_position_x = original_item_position[0]
                original_item_position_y = original_item_position[1]
                dpg.set_item_pos(child, [original_item_position_x, original_item_position_y - 50])

    #The window where the user can select an object type, and create a new object of that type
    with dpg.window(tag="SaveObj", label=tab + "Create Obj", pos=(300,200), no_collapse=True, width=875, height=600) as save_obj:
        dpg.bind_item_theme(save_obj, window_theme)#assigns the window theme to this window 
        with dpg.menu_bar():
            dpg.add_text("Save Object")
            dpg.add_button(label="Close Window",
            tag="CloseWindow",
            pos=(dpg.get_viewport_width() - 175,0),
            width=150)
            dpg.bind_item_theme("CloseWindow", menu_bar_drop_down_theme)
            
        #Creates the text preceding the drop down menu
        dpg.add_button(
        label="Object Type:",
        tag="ChooseObjTypeTxt", 
        width=185, 
        height=31.775, 
        parent="GetKeyValues",
        pos=(offset + 35, 75))
        
        #Gets a list of every object type name 
        obj_type_names = GetObjTypeNames()
        
        #Creates a drop down menu populated by the name of every object type, when an object type is chosen, 
        #a series of value entry fields are created that pertain to the keys/attributes of that object type
        choose_obj_type = dpg.add_combo(
        items=obj_type_names, 
        default_value=obj_type_names[0],
        callback=GenObjValueFields,
        tag="ChooseObjType",
        width=600,
        pos=(offset + 225, 75))
        
        #if the user has modified window's shape, retain their inputted information
        if saved_input == None:
            pass
        else:
            for user_input in saved_input:
                if user_input[0] == "ChooseObjType":
                    dpg.set_value("ChooseObjType", user_input[1])
                    if len(saved_input) == 1:
                        pass
                    else:
                        GenObjValueFields("null", user_input[1], "null")
            
        with dpg.item_handler_registry() as auto_size_window:
            dpg.add_item_hover_handler(callback=AutoSizeWindow)
        dpg.bind_item_handler_registry("SaveObj", auto_size_window)
        
    
    if primary_window == "SaveObj":   
        dpg.set_primary_window("SaveObj", True)
    
    dpg.set_primary_window("PrimaryWindow", False)
    dpg.delete_item("PrimaryWindow")
    dpg.delete_item("SaveObj")
    primary_window = "SaveObj"
    CreateObjWindow(calc_offset=True, saved_input=saved_input)
    
def NewObj_Window():
    global primary_window
    with dpg.window(label="NewObj", tag="NewObj"):
        dpg.bind_item_theme("NewObj", window_theme)
        
        #Intialize the layout table 
        table_height_percentage = .2 #1.0 is maximum viewport height
        table_width_percentage = .2  #1.0 is maximum viewport width
        column_amount = 3
        row_amount = 1
        CreateTable("FormatTable", "NewObj", column_amount=column_amount, row_amount=row_amount)
                   
    #Store and close the previous window, then set this window to the primary window 
    previous_window = primary_window 
    dpg.delete_item(previous_window)
    dpg.set_primary_window("NewObj", True)
    primary_window = "NewObj"
    
    #Configure the table's width and height according to the viewport's dimensions
    viewport_width = dpg.get_viewport_width()
    viewport_height = dpg.get_viewport_height()
    table_width = viewport_width - (table_width_percentage*(viewport_width))
    table_height = viewport_height - (table_height_percentage*(viewport_height))
    dpg.configure_item("FormatTable", height=table_height, width=table_width)
    dpg.set_item_pos("TablePos", [(viewport_width/2) - (table_width/2), (viewport_height - abs((table_height_percentage/2)*(viewport_height) - viewport_height))])
    
    #Resize the rows such that they are all of equal height 
    ResizeRows("FormatTable")

    max_column_width = -1 #-1 stretches a widget to fill all avaliable space 
    default_widget_height = 31.775
    column_width_weights = [.75,2.0,.25]
    #Force all width weights to be rational, a column should always receive a fraction of the tables width, and the width weights should always add up to the number of columns 
    if sum(column_width_weights) > column_amount or sum(column_width_weights) < column_amount:
        column_index = 0
        for column_width_weight in column_width_weights:
            column_width_weights[column_index] = 1.0
            column_index+=1
        print("error: sum of column width weights exceeded or fell below the available table width, thus all column widths have been defaulted to equal")
    else:
        pass
    for column_width_weight in column_width_weights:
        if column_width_weight <= 0:
            column_index = 0
            for column_width_weight in column_width_weights:
                column_width_weights[column_index] = 1.0
                column_index+=1
            print("error: a column width weight was less than or equal to 0, thus all column widths have been defaulted to equal")
        else:
            pass
    
    #Assign Column Width Weights
    column_index=1
    while column_index <= column_amount:
        dpg.configure_item("Col" + str(column_index), init_width_or_weight=column_width_weights[column_index-1])
        column_index+=1 
        
    #Create the object type selection menu, and the ability to generate value input fields 
    def GenSubValueFields(sender, app_data, user_data):
        row = dpg.get_item_alias(dpg.get_item_parent(sender))
        print(row)
        insert_pos = int(row[row.find("Row") + len("Row"):row.find("C")])
        InsertRow("FormatTable", insert_pos=insert_pos, row_amount=1)

    def GenValueFields(sender, app_data, user_data):
        #Generate value entry fields for each of the object type's attributes upon selection 

        #Check and delete any pre-existing input fields
        if dpg.does_item_exist("Row2") == True:
            dpg.delete_item("GetKeyValues")  
            row_index = 2
            while dpg.does_item_exist("Row" + str(row_index)) == True:
                row_index += 1
            last_row = row_index
            row_index = 2
            while row_index <= last_row:
                dpg.delete_item("Row" + str(row_index))
                row_index+=1 
                
        #Get the keys/attributes of the object type
        obj_type = app_data
        obj_type_keys = GetObjKeys("null", "null", obj_type)
        
        #iterate through the keys and apply formatting to their titles   
        formatted_keys = []
        key_ID = 1
        for key in obj_type_keys:
            if "_" in key:#There is more than one word in the key, format accordingly 
                formatted_key = [] 
                last_seperator_pos = 0
                seperator_index = 0
                while seperator_index <= key.count("_"):
                    if seperator_index == key.count("_"):
                        seperator_pos = len(key)
                        last_seperator_pos += 1
                    else:
                        seperator_pos = key.find("_", last_seperator_pos)
                    formatted_key.append(key[last_seperator_pos:seperator_pos])
                    last_seperator_pos = seperator_pos
                    seperator_index += 1
                key_part_index = 0
                for key_part in formatted_key:
                    key_part = key_part[0].upper() + key_part[1:len(key_part)]
                    formatted_key[key_part_index] = key_part
                    key_part_index += 1
                key = ""
                for key_part in formatted_key:
                    if key_part != formatted_key[len(formatted_key) - 1]:
                        key = key + key_part + " "
                    else:
                        key = key + key_part + ":"
                formatted_keys.append(key)
            else:#There is only one word in the key, format accordingly 
                key = key[0].upper() + key[1:len(key)] + ":"
                formatted_keys.append(key)
        
        #Intialize the rows 
        row_amount = len(formatted_keys)
        AppendCells("FormatTable", column_amount=0, row_amount=row_amount)
        
        #Insert the value inputs and value labels into the cells 
        row_index = 2
        widget_index = 1
        for key in formatted_keys:
            dpg.add_button(label=key, tag="InputTxt" + str(widget_index), width=max_column_width, parent="Row" + str(row_index) + "Col1" + "Content")
            dpg.add_input_text(tag="Input" + str(widget_index), width=max_column_width, parent="Row" + str(row_index) + "Col2" + "Content")
            dpg.add_button(label="Add Value", tag="AddInput" + str(widget_index), width=max_column_width, parent="Row" + str(row_index) + "Col3" + "Content", callback=GenSubValueFields, user_data=widget_index)
            dpg.bind_item_theme("Input" + str(widget_index), text_input_theme)
            dpg.bind_item_theme("InputTxt" + str(widget_index), text_box_theme)
            dpg.bind_item_theme("AddInput" + str(widget_index), add_theme)
            row_index += 1
            widget_index += 1
            
        
            
    obj_type_names = GetObjTypeNames()
    obj_type_select_txt = dpg.add_button(label="Object Type:", tag="ObjTypeSelectTxt", parent="Row1Col1Content", height=default_widget_height, width=max_column_width)
    dpg.bind_item_theme("ObjTypeSelectTxt", text_box_theme)
    obj_type_select = dpg.add_combo(tag="ObjTypeSelect", parent="Row1Col2Content", items=obj_type_names, default_value=obj_type_names[0], width=max_column_width, callback=GenValueFields)
    dpg.bind_item_theme("ObjTypeSelect", combo_box_theme)



#GUI Code
dpg.create_context() 
dpg.create_viewport(title='Drocula', width=600, height=600)


#Fonts
with dpg.font_registry():
    default_font = dpg.add_font("PixelPurl.ttf", 25)
  
dpg.bind_font(default_font)

#Themes
with dpg.theme() as window_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_MenuBarBg, (latte[0], latte[1], latte[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_WindowBg, (white[0], white[1], white[2]), category=dpg.mvThemeCat_Core)
    
with dpg.theme() as global_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_WindowBg, (blue[0], blue[1], blue[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_TitleBg, (light_gray[0], light_gray[1], light_gray[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_TitleBgActive, (light_gray[0], light_gray[1], light_gray[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_TitleBgCollapsed, (light_gray[0], light_gray[1], light_gray[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_ButtonHovered, (light_blue[0], light_blue[1], light_blue[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_Text, (black[0], black[1], black[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_Border, (dark_gray[0], dark_gray[1], dark_gray[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_Button, (yellow[0], yellow[1], yellow[2]), category=dpg.mvThemeCat_Core) 
        dpg.add_theme_style(dpg.mvStyleVar_FrameBorderSize, 2, category=dpg.mvThemeCat_Core)
        dpg.add_theme_style(dpg.mvStyleVar_WindowBorderSize, 2, category=dpg.mvThemeCat_Core)

with dpg.theme() as text_box_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_Button, (apricot[0], apricot[1], apricot[2]), category=dpg.mvThemeCat_Core)
with dpg.theme() as combo_box_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_Button, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core) 
        dpg.add_theme_color(dpg.mvThemeCol_ButtonHovered, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core)
        dpg.add_theme_color(dpg.mvThemeCol_FrameBg, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core) #drop down menu button backround
        dpg.add_theme_color(dpg.mvThemeCol_Header, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core) #drop down menu item color
        dpg.add_theme_color(dpg.mvThemeCol_HeaderHovered, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core) #drop down menu item color
        dpg.add_theme_color(dpg.mvThemeCol_HeaderActive, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core) #drop down menu item color
        dpg.add_theme_color(dpg.mvThemeCol_PopupBg , (light_cream[0], light_cream[1], light_cream[2]), category=dpg.mvThemeCat_Core) #drop down menu background color
with dpg.theme() as text_input_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_FrameBg, (heavy_custard[0], heavy_custard[1], heavy_custard[2]), category=dpg.mvThemeCat_Core) #drop down menu button backround
with dpg.theme() as table_background_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_ChildBg, (white[0], white[1], white[2]), category=dpg.mvThemeCat_Core) #drop down menu button backround
with dpg.theme() as add_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_Button, (sea_foam[0], sea_foam[1], sea_foam[2]), category=dpg.mvThemeCat_Core) 
with dpg.theme() as del_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_Button, (coral[0], coral[1], coral[2]), category=dpg.mvThemeCat_Core) 
dpg.bind_theme(global_theme)   








DefaultWindow()
   





   







dpg.setup_dearpygui()
dpg.show_viewport()
dpg.maximize_viewport()


dpg.set_primary_window("DefaultWindow", True)

dpg.start_dearpygui()
dpg.destroy_context()
